import os
from django.shortcuts import get_object_or_404
from pandas.tseries.frequencies import _is_multiple
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.authentication import JWTAuthentication
from rest_framework.decorators import api_view
from rest_framework.views import APIView
from interview.models import Interview,InterviewCandidateStatus
from interview.serializer import InterviewListSerializer,InterviewCandidateStatusCareerSerializer
from job.decorators import credit_limit_email_notifier, deduct_credit_decorator
from settings.decorators import check_feature_availability, check_subscription_and_credits_for_ai, handle_ai_credits, handle_career_site_ai_credits
from .import helper
from common.utils import makeResponse
from common.mail_utils import CandidateCreateMailer
from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.authentication import JWTAuthentication
from rest_framework import mixins, generics
from .models import *
from rest_framework import status
from django.http import JsonResponse
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.authentication import JWTAuthentication
from .serializer import *
import json
from rest_framework import viewsets, status
from job.models import *
from settings.models import *
from django.http import Http404, JsonResponse, StreamingHttpResponse

def serialize_usage_metadata(obj):
    """Recursively serialize objects to JSON-compatible format"""
    if obj is None:
        return None
    if isinstance(obj, (int, float, str, bool)):
        return obj
    if hasattr(obj, '__dict__'):
        return {k: serialize_usage_metadata(v) for k, v in vars(obj).items()}
    if isinstance(obj, (list, tuple)):
        return [serialize_usage_metadata(item) for item in obj]
    if isinstance(obj, dict):
        return {k: serialize_usage_metadata(v) for k, v in obj.items()}
    # For other objects, try to convert to string
    try:
        return str(obj)
    except:
        return None
import pandas as pd
from django.conf import settings
import os, io
from pdfminer.high_level import extract_text
from io import BytesIO
from openai import OpenAI
from docx import Document
from django.db.models import Q, Count
from django.utils import timezone
from datetime import datetime, timedelta
import re
from .assessmentchacke import check_assessment
import threading
from .resume_mapper import map_resume_to_module
from .candidate_evaluation import generate_rjms
from job.views import SmartJobSearchAPIView
import threading
import logging
import imaplib
import email
from concurrent.futures import ThreadPoolExecutor
from rest_framework.response import Response
from django.http import JsonResponse
from rest_framework.views import APIView
import logging
from django.db import transaction
from django.utils import timezone
from django.core.mail import send_mail
import openai
from django.db.models import Q
# import google.genai as genai

# API_KEY = "..." removed for security
client = OpenAI(api_key=settings.OPENAI_API_KEY)

class ApplyApi(APIView):

    def post(self, request):
        data = helper.applyJob(request)
        return makeResponse(data)

class ApplyJobApi(APIView):

    def post(self, request):
        data = helper.applyWebformJob(request)
        return makeResponse(data)        

class ApplicationsApi(APIView):

    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        data = helper.getApplications(request)
        return makeResponse(data)

    def post(self, request):
        data = helper.updateApplication(request)
        return makeResponse(data)

    def delete(self, request):
        data = helper.deleteApplication(request)
        return makeResponse(data)    

class CandidatesApi(APIView):

    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        data = helper.getCandidates(request)
        return makeResponse(data)


class CreateCandidateUsingResume(APIView):

    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = helper.createCandidate(request)
        return makeResponse(data)

        

class ApplicationsResumeApi(APIView):

    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = helper.updateResume(request)
        return makeResponse(data)  

class ApplicationsResumeParseApi(APIView):

    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = helper.parseResume(request)
        return makeResponse(data)                 

class CandidateDetailsApi(APIView):

    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        data = helper.candidateDetails(request)
        return makeResponse(data)                

class NoteApi(APIView):

    # authentication_classes = [JWTAuthentication]
    # permission_classes = [IsAuthenticated]

    def get(self, request):
        data = helper.getAllNotes(request)
        return makeResponse(data)

    def post(self, request):
        data = helper.saveNote(request)
        return makeResponse(data)

    def delete(self, request):
        data = helper.deleteNote(request)
        return makeResponse(data)    
    
    def put(self, request):
        data = helper.updateNote(request)
        return makeResponse(data) 

class DetailNoteApi(mixins.RetrieveModelMixin, generics.GenericAPIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    queryset = Note.objects.all()
    serializer_class = NoteSerializer

    def get(self, request, *args, **kwargs):
        return self.retrieve(request, *args, **kwargs)


# Update Notes view 
class NotesUpdateApi(
    generics.UpdateAPIView
    ):
    queryset = Note.objects.all()
    serializer_class = NoteSerializer

    def perform_update(self, serializer):
        instance = serializer.save()

class CreateCandidateUsingWebForm(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = helper.createCandidatewithoutResumeParser(request)
        return makeResponse(data) 

class UpdateCandidatePipelineStatus(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = helper.updatePipelineStatus(request)
        return makeResponse(data) 

# Details ApplicationWebForm view
class ApplicationWebFormByJobApi(
    generics.ListAPIView
    ):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        lookup_field = self.kwargs['lookup_field']
        lookup_value = self.kwargs['lookup_value']
        queryset = ApplicantWebForm.objects.filter(**{lookup_field: lookup_value})
        return queryset

    serializer_class = ApplicantWebFormSerializer

# Update ApplicationWebForm view 
class ApplicationWebFormUpdateApi(
    generics.UpdateAPIView
    ):
    queryset = ApplicantWebForm.objects.all()
    serializer_class = ApplicantWebFormSerializer

    def perform_update(self, serializer):
        instance = serializer.save()

# Update ApplicationWebForm view 
class ApplicationWebFormUpdateApi(
    generics.UpdateAPIView
    ):
    queryset = ApplicantWebForm.objects.all()
    serializer_class = ApplicantWebFormSerializer

    def perform_update(self, serializer):
        instance = serializer.save()
# Delete ApplicationWebForm view 
class ApplicationWebFormDeleteApi(
    generics.DestroyAPIView
    ):
    queryset = ApplicantWebForm.objects.all()
    serializer_class = ApplicantWebFormSerializer

    def perform_destroy(self, instance):
        super().perform_destroy(instance)
 
# Create ApplicationWebForm view
class ApplicationWebFormCreateApi(APIView):
    def post(self, request):
        data = helper.saveApplicantWebForms(request)
        return makeResponse(data)

class AssignJob(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        data = helper.assignJob(request)
        return makeResponse(data)

class CandidateStats(APIView):

    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        data = helper.getCandidateStats(request)
        return makeResponse(data)

class HiringFunnelApi(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        data = helper.get_hiring_funnel(request)
        return makeResponse(data)

class HiresBySourceApi(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        data = helper.get_hires_by_source(request)
        return makeResponse(data)

class DashboardCardsApi(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        data = helper.get_dashboard_cards(request)
        return makeResponse(data)
   
class UpdateCandidateStatusApi(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        data = helper.updateCandidatePipelineStage(request)
        return makeResponse(data)
    
class CandidateJsonDataView(APIView):
    def get(self,request):
        data = Candidate.objects.all()
        serializer = CandidateJsonDataSerializer(data,many=True)
        return Response(serializer.data)

def get_safe_performed_by(request):
    """
    Safely get the performed_by user from request.
    Returns None if user is not authenticated or is AnonymousUser.
    """
    if request.user and hasattr(request.user, 'account_id') and not request.user.is_anonymous:
        return request.user
    return None   

class CandidateApiView(APIView):
    # authentication_classes = [JWTAuthentication]
    # permission_classes = [IsAuthenticated]

    def get(self, request,  pk=None,*args, **kwargs):
        if pk is not None:
            # Retrieve a specific candidate by ID
            candidate = self.get_object(pk)
            serializer = CandidateSerializer(candidate,context={'request': self.request})
            return Response(serializer.data)
        else:
            # Retrieve all candidates
            # candidates = Candidate.objects.filter(job__company__pk=self.request.user.company_id)
            candidates = Candidate.objects.filter(company__pk=self.request.user.company_id).distinct()
            serializer = CandidateSerializer(candidates, many=True,context={'request': self.request})
            return Response(serializer.data)

    def post(self, request, *args, **kwargs):
        data = request.data.copy()
                
        job = data.get('job', None)
        webform_candidate_data = data.get('webform_candidate_data', {})
        
        if isinstance(webform_candidate_data, str):
            try:
                webform_candidate_data = json.loads(webform_candidate_data)
            except json.JSONDecodeError:
                return Response({"error": "Invalid webform_candidate_data format"}, status=status.HTTP_400_BAD_REQUEST)
    
        # Get email from the parsed data
        email = webform_candidate_data.get('Personal Details', {}).get('email')
        
        # if not job or not email:
        #     return Response({"error": "Job and email are required"}, status=status.HTTP_400_BAD_REQUEST)
        if job:   
            if 'job' not in webform_candidate_data['Professional Details']:
                webform_candidate_data['Professional Details']['job'] = [int(job)]
                # Update the data with modified webform_candidate_data (convert back to JSON string)
                data['webform_candidate_data'] = json.dumps(webform_candidate_data)
        
        # Get the company ID from the job
        # try:
        #     job_obj = Job.objects.get(id=job)  # Assuming job is a list with one ID
        #     company_id = job_obj.company.id
        # except Job.DoesNotExist:
        #     return Response({"error": "Job not found"}, status=status.HTTP_404_NOT_FOUND)
        company_id = None
        existing_candidate = None
        job_obj = None
        if job:
            job_obj = Job.objects.get(id=job)  # Assuming job is a list with one ID
            company_id = job_obj.company.id

        
        # Check if candidate exists with same email and company
        if company_id:
            existing_candidate = Candidate.objects.filter(
                job__company__id=company_id
            ).filter(
                webform_candidate_data__contains={'Personal Details': {'email': email}}
            ).first()
        
        if existing_candidate:
            # Add new job to existing candidate 
            existing_candidate.job.add(job_obj)
            
            current_webform_data = existing_candidate.webform_candidate_data
            if 'Professional Details' not in current_webform_data:
                current_webform_data['Professional Details'] = {}
            
            # Get existing jobs or initialize empty list
            current_jobs = current_webform_data['Professional Details'].get('job', [])
            if not isinstance(current_jobs, list):
                current_jobs = []
                
            # Convert all job IDs to integers
            current_jobs = [int(job_id) for job_id in current_jobs if str(job_id).isdigit()]
            
            # Add new job if not already present
            new_job_id = int(job)
            if new_job_id not in current_jobs:
                current_jobs.append(new_job_id)

            # Update the jobs list
            current_webform_data['Professional Details']['job'] = current_jobs
            
            performed_by = get_safe_performed_by(request)
            CandidateTimeline.log_activity(existing_candidate, 'CANDIDATE_CREATED', 'Candidate Created',description=f"Added new jobs", job=new_job_id, performed_by=performed_by,pipeline_stage="Screening")
            JobTimeline.log_activity(job=job_obj,candidate=existing_candidate, activity_type= 'CANDIDATE_JOB_APPLIED', title= 'Candidate Job Applied',description=f"Candidate job applied", performed_by=performed_by)
            # Save the updated webform_candidate_data
            existing_candidate.webform_candidate_data = current_webform_data
            existing_candidate.save()
            serializer = CandidateSerializer(existing_candidate)
            
            if serializer.instance.source == 'career site':
                try:
                    email_Settings = EmailSettings.objects.get(user_id=job_obj.created_by, company=job_obj.company)
                    if email_Settings:
                        message = f"Thank you for applying to {job_obj.title} at {job_obj.company.name}. We have received your application and will review it shortly."
                        sendMail = CandidateCreateMailer(
                            emails=serializer.instance.webform_candidate_data["Personal Details"]["email"],
                            message=message,
                            email_Settings=email_Settings,
                            candidate=serializer.instance,
                            job=job_obj,
                            account=job_obj.created_by
                        )
                        sendMail.start()
                except (EmailSettings.DoesNotExist, EmailSettings.MultipleObjectsReturned):
                    # Log that email settings not found for this user
                    pass                      
            
            return Response(serializer.data, status=status.HTTP_200_OK)
        
        # Add company only when there's no job or if job is provided
        if job:
            data['company'] = company_id
        elif not job:
            data['company'] = request.user.company_id
        
        # Create new candidate if none exists
        serializer = CandidateSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            candidate = serializer.instance

            if job:
                try:
                    new_job_id = int(job)
                except (TypeError, ValueError):
                    new_job_id = None
            else:
                new_job_id = None
            
            if serializer.instance.source == 'career site':
                try:
                    email_Settings = EmailSettings.objects.get(user_id=job_obj.created_by, company=job_obj.company)
                    if email_Settings:
                        message = f"Thank you for applying to {job_obj.title} at {job_obj.company.name}. We have received your application and will review it shortly."
                        sendMail = CandidateCreateMailer(
                            emails=serializer.instance.webform_candidate_data["Personal Details"]["email"],
                            message=message,
                            email_Settings=email_Settings,
                            candidate=serializer.instance,
                            job=job_obj,
                            account=job_obj.created_by
                        )
                        sendMail.start()
                except (EmailSettings.DoesNotExist, EmailSettings.MultipleObjectsReturned):
                    # Log that email settings not found for this user
                    pass
                except EmailCategory.DoesNotExist:
                    # Log that email category not found
                    pass
                except Exception as e:
                    # Log any other unexpected errors
                    pass
            
            performed_by = get_safe_performed_by(request)
            if job:
                CandidateTimeline.log_activity(serializer.instance, 'CANDIDATE_CREATED', 'Candidate Created', job=data.get('job'), performed_by=performed_by,pipeline_stage="Screening")
            else:
                CandidateTimeline.log_activity(serializer.instance, 'CANDIDATE_CREATED', 'Candidate Created', performed_by=performed_by)
            if job_obj:    
                JobTimeline.log_activity(job=job_obj,candidate=serializer.instance, activity_type= 'CANDIDATE_JOB_APPLIED', title= 'Candidate Job Applied',description=f"Candidate job applied", performed_by=performed_by)

            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def put(self, request, pk, *args, **kwargs):
        # Update a candidate
        candidate = self.get_object(pk)
        old_job_ids = list(candidate.job.all().values_list('id', flat=True))
        
        data = request.data.copy()
        job_ids = data.pop('job', None)

        serializer = CandidateSerializer(candidate, data, partial=True)
        updated_fields = {}
        
        # Handle webform_candidate_data job updates
        webform_data = request.data.get('webform_candidate_data', {})
        if isinstance(webform_data, str):
            try:
                webform_data = json.loads(webform_data)
            except json.JSONDecodeError:
                return Response({'error': 'Invalid JSON in webform_candidate_data'}, status=400)
        
        jobids=webform_data.get('Professional Details', {}).get('job', [])
        jobs_ids=[]
        performed_by = get_safe_performed_by(request)
        if len(jobids) > len(old_job_ids):
            for jobid in jobids:
                if jobid not in old_job_ids:
                    jobs_ids.append(jobid)
            CandidateTimeline.log_activity(candidate, 'CANDIDATE_UPDATED', 'Candidate Updated',description=f"Added new jobs", job=jobs_ids, performed_by=performed_by)
        elif len(jobids) < len(old_job_ids):
            for jobid in old_job_ids:
                if jobid not in jobids:
                    jobs_ids.append(jobid)
            CandidateTimeline.log_activity(candidate, 'CANDIDATE_UPDATED', 'Candidate Updated',description=f"Removed jobs", job=jobs_ids, performed_by=performed_by)
        else:
            # Correctly identify added and removed jobs
            added_jobs = [jobid for jobid in jobids if jobid not in old_job_ids]
            removed_jobs = [jobid for jobid in old_job_ids if jobid not in jobids]
            
            if added_jobs:
                CandidateTimeline.log_activity(candidate, 'CANDIDATE_UPDATED', 'Candidate Updated',description=f"Added jobs", job=added_jobs, performed_by=performed_by)
            if removed_jobs:
                CandidateTimeline.log_activity(candidate, 'CANDIDATE_UPDATED', 'Candidate Updated',description=f"Removed jobs", job=removed_jobs, performed_by=performed_by)
            if not added_jobs and not removed_jobs:
                CandidateTimeline.log_activity(candidate, 'CANDIDATE_UPDATED', 'Candidate Updated',description=f"Candidate Updated", performed_by=performed_by)
        
        
            
        if serializer.is_valid():
            serializer.save()
            
            if job_ids is not None:
                jobs = Job.objects.filter(id__in=job_ids)  # Filter jobs by the received IDs
                candidate.job.set(jobs)

            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk, *args, **kwargs):
        # Delete a candidate
        candidate = self.get_object(pk)
        candidate.delete()
        return Response({
            'code': 200,
            'msg': 'Candidate deleted successfully'
        })

    def get_object(self, pk):
        # Helper method to get a candidate by pk
        try:
            return Candidate.objects.get(pk=pk)
        except Candidate.DoesNotExist:
            raise Http404

logger = logging.getLogger(__name__)
logger = logging.getLogger(__name__)

def run_assessment_background(candidate_id: int, assessment_data: dict, job_id: int, resume_data: str) -> None:
    """
    Process assessment data in background and trigger RJMS analysis.
    
    Args:
        candidate_id: ID of the candidate
        assessment_data: Assessment form data
        job_id: ID of the job
        resume_data: Resume text data
    """
    logger.info('Process assessment data in background')
    try:
        with transaction.atomic():
            candidate = Candidate.objects.select_for_update().get(id=candidate_id)
            job = Job.objects.get(id=job_id)
            
            # Check if job has assessment configuration
            if not hasattr(job, 'assesment') or job.assesment is None:
                logger.info(f"No assessment for job {job_id}, running RJMS without assessment")
                assessment_result = None
                sms = 0
            else:
                assessment_config = job.assesment
                result = check_assessment(assessment_data, assessment_config.form, company=job.company, job_id=job.id)
                assessment_result = {
                    "assessment_name": assessment_config.name,
                    "questions": result,
                }
                sms = result[-1]["weighted_percentage"]
                
                # Update candidate's assessment data
                current_data = candidate.assessment_data or []
                if not isinstance(current_data, list):
                    current_data = [current_data] if current_data else []
                current_data.append(assessment_result)
                candidate.assessment_data = current_data
                candidate.save()
            
        # Start RJMS analysis in background (with or without assessment)
        rjms_thread = threading.Thread(
            target=run_rjms_background,
            args=(candidate_id, job.id, resume_data, assessment_result, sms),
            daemon=True,
            name=f"RJMS-Analysis-{candidate_id}-{job.id}"
        )
        rjms_thread.start()
        
    except Exception as e:
        logger.error(
            "Error in run_assessment_background: %s",
            str(e),
            exc_info=True,
            extra={
                'candidate_id': candidate_id,
                'job_id': job_id,
                'assessment_data_keys': list(assessment_data.keys()) if isinstance(assessment_data, dict) and assessment_data else None
            }
        )

def run_rjms_background(candidate_id: int, job_id: int, resume_text: str, assessment_result: dict, sms:float) -> None:
    """
    Run RJMS analysis in background and save results.
    
    Args:
        candidate_id: ID of the candidate
        job_id: ID of the job
        resume_text: Text content of the resume
        assessment_result: Assessment results to include in analysis
    """
    try:
        with transaction.atomic():
            candidate = Candidate.objects.select_for_update().get(id=candidate_id)
            job = Job.objects.select_related('company').get(id=job_id)
            
            # Generate RJMS analysis
            result = generate_rjms(
                resume=resume_text,
                job_id=job_id,
                company=job.company,
                assessment_data=assessment_result,
                sms=sms
            )
            
            if isinstance(result, dict):
                # Update or create RJMS analysis record with only valid fields
                sms_score = result.get('sms', 0)  # Extract sms_score from result if available
                
                RJMSAnalysis.objects.update_or_create(
                    candidate=candidate,
                    job=job,
                    defaults={
                        'result': result,
                        # 'sms_score': sms_score,
                    }
                )
    except Exception as e:
        logger.error(
            "Error in run_rjms_background: %s",
            str(e),
            exc_info=True,
            extra={
                'candidate_id': candidate_id,
                'job_id': job_id,
                'assessment_name': assessment_result.get('assessment_name')
            }
        )

class CareerSiteCandidateApiView(APIView):

    @deduct_credit_decorator('candidate_apply')
    @deduct_credit_decorator('Candidate_Evaluation')
    @deduct_credit_decorator('Resume_Parcing_using_AI')
    @credit_limit_email_notifier('Candidate_Evaluation')
    @credit_limit_email_notifier('Resume_Parcing_using_AI')
    # @handle_career_site_ai_credits(usage_codes=["Resume_Parcing_using_AI", "Rjms", "assessment_check"])
    def post(self, request, *args, **kwargs):
        data = request.data.copy()
        job = data.get('job', None)
        webform_candidate_data = data.get('webform_candidate_data', {})
        assessment_data = request.data.get("assessment_data")
        # Use a separate variable for the job's assessment configuration
        # assessment_config = None
        # assessment_result = None
        # if job:
        #     try:
        #         assessment_config = Job.objects.get(id=job).assesment
        #     except Job.DoesNotExist:
        #         assessment_config = None
        # If assessment data is provided, attempt to score it now
        # if assessment_data and assessment_config:
        #     assessment_result = check_assessment(assessment_data, assessment_config.form)
        #     assessment_result = {"assessment_name":assessment_config.name,"questions":assessment_result}
        # Ensure the computed result is placed into the payload for serializer
        # For multipart/form-data, JSONField expects a valid JSON string
        # if assessment_result is not None:
        #     # Store as a simple list of attempts: [attempt0, attempt1, ...]
        #     try:
        #         data['assessment_data'] = json.dumps([assessment_result])
        #     except Exception:
        #         data['assessment_data'] = [assessment_result]
        if isinstance(webform_candidate_data, str):
            try:
                webform_candidate_data = json.loads(webform_candidate_data)
            except json.JSONDecodeError:
                return Response({"error": "Invalid webform_candidate_data format"}, status=status.HTTP_400_BAD_REQUEST)
    
        # Get email from the parsed data
        email = webform_candidate_data.get('Personal Details', {}).get('email')
        
        
        if job:   
            if 'job' not in webform_candidate_data['Professional Details']:
                webform_candidate_data['Professional Details']['job'] = [int(job)]
                # Update the data with modified webform_candidate_data (convert back to JSON string)
                data['webform_candidate_data'] = json.dumps(webform_candidate_data)
        
        company_id = None
        existing_candidate = None
        job_obj = None
        if job:
            job_obj = Job.objects.get(id=job)  # Assuming job is a list with one ID
            company_id = job_obj.company.id

        credit_wallets = CreditWallet.objects.filter(feature__code__in=["Resume_Parcing_using_AI",'Candidate_Evaluation'], company=company_id)
        
        # Check if candidate exists with same email and company
        if company_id:
            existing_candidate = Candidate.objects.filter(
                job__company__id=company_id
            ).filter(
                webform_candidate_data__contains={'Personal Details': {'email': email}}
            ).first()
            
        
        if existing_candidate:
            # Add new job to existing candidate 
            existing_candidate.job.add(job_obj)
            try:
                if request.user and hasattr(request.user, 'account_id') and not request.user.is_anonymous:
                    existing_candidate.account = request.user
                    existing_candidate.save()  # Don't forget to save the changes
            except Exception as e:
                # Log the error for debugging
                logger.warning(f"Error updating candidate account: {str(e)}")
                # Return an error response instead of continue
                
            
            current_webform_data = existing_candidate.webform_candidate_data
            if 'Professional Details' not in current_webform_data:
                current_webform_data['Professional Details'] = {}
            
            # Get existing jobs or initialize empty list
            current_jobs = current_webform_data['Professional Details'].get('job', [])
            if not isinstance(current_jobs, list):
                current_jobs = []
                
            # Convert all job IDs to integers
            current_jobs = [int(job_id) for job_id in current_jobs if str(job_id).isdigit()]
            
            # Add new job if not already present
            new_job_id = int(job)
            if new_job_id not in current_jobs:
                current_jobs.append(new_job_id)

            # Update the jobs list
            current_webform_data['Professional Details']['job'] = current_jobs
            
            performed_by = get_safe_performed_by(request)
            CandidateTimeline.log_activity(existing_candidate, 'CANDIDATE_CREATED', 'Candidate Created',description=f"Added new jobs", job=new_job_id, performed_by=performed_by,pipeline_stage="Screening")
            JobTimeline.log_activity(job=job_obj,candidate=existing_candidate, activity_type= 'CANDIDATE_JOB_APPLIED', title= 'Candidate Job Applied',description=f"Candidate job applied", performed_by=performed_by)
            # Save the updated webform_candidate_data
            existing_candidate.webform_candidate_data = current_webform_data
            # Merge assessment_result into existing candidate's assessment_data as a simple list of attempts
            # if assessment_result is not None:
            #     try:
            #         current_data = existing_candidate.assessment_data
            #         # Empty -> start a new attempts list
            #         if current_data in (None, {}, []):
            #             existing_candidate.assessment_data = [assessment_result]
            #         elif isinstance(current_data, list):
            #             # If it's a flat list of question dicts (single attempt), wrap first
            #             if len(current_data) > 0 and isinstance(current_data[0], dict):
            #                 existing_candidate.assessment_data = [current_data]
            #             else:
            #                 existing_candidate.assessment_data = current_data
            #             existing_candidate.assessment_data.append(assessment_result)
            #         elif isinstance(current_data, dict):
            #             # Migrate from job-keyed map to flat attempts list
            #             attempts = []
            #             for v in current_data.values():
            #                 if isinstance(v, list):
            #                     # v may be a single attempt (list of question dicts) or attempts
            #                     if len(v) > 0 and isinstance(v[0], dict):
            #                         attempts.append(v)
            #                     else:
            #                         attempts.extend(v)
            #             attempts.append(assessment_result)
            #             existing_candidate.assessment_data = attempts if attempts else [assessment_result]
            #         else:
            #             # Unknown shape -> coerce
            #             existing_candidate.assessment_data = [current_data, assessment_result]
            #     except Exception:
            #         existing_candidate.assessment_data = [assessment_result]
            existing_candidate.save()

            serializer = CandidateSerializer(existing_candidate)

            assessment_thread = threading.Thread(
                    target=run_assessment_background,
                    args=(
                        existing_candidate.id,
                        assessment_data,
                        job.id if isinstance(job, Job) else job,  # Handle both object and ID
                        data.get('resume_data', '')
                    ),
                    daemon=True,
                    name=f"Assessment-{existing_candidate.id}-{getattr(job, 'id', job)}"
                )
            assessment_thread.start()
            
            if data.get('source', None) != '':
                try:
                    email_Settings = EmailSettings.objects.get(user_id=job_obj.created_by, company=job_obj.company)
                    if email_Settings:
                        message = f"Thank you for applying to {job_obj.title} at {job_obj.company.name}. We have received your application and will review it shortly."
                        sendMail = CandidateCreateMailer(
                            emails=serializer.instance.webform_candidate_data["Personal Details"]["email"],
                            message=message,
                            email_Settings=email_Settings,
                            candidate=serializer.instance,
                            job=job_obj,
                            account=job_obj.created_by
                        )
                        sendMail.start()
                except (EmailSettings.DoesNotExist, EmailSettings.MultipleObjectsReturned):
                    # Log that email settings not found for this user
                    pass
            
            return Response(serializer.data, status=status.HTTP_200_OK)
        
        # Add company only when there's no job or if job is provided
        if job:
            data['company'] = company_id
        elif not job:
            data['company'] = request.user.company_id

        if 'assessment_data' in data:
            data.pop('assessment_data')
        
        # Create new candidate if none exists
        # Set the account from the job's created_by if it's a career site submission
        if 'source' in data and data['source'] != '':
            data['account'] = request.user.id
        # Fall back to request user if available
        elif hasattr(request, 'user') and request.user.is_authenticated:
            data['account'] = request.user.id
        
        block_index = 0
        if credit_wallets:
            cw_for_block = credit_wallets.first()
            if cw_for_block and cw_for_block.total_credit > 0:
                block_index = cw_for_block.used_credit // cw_for_block.total_credit

        # today + block_index days
        updated_date = datetime.now().date() + timedelta(days=block_index)

        data['updated'] = updated_date

        serializer = CandidateSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            candidate = serializer.instance

            resume_path = data.get("resume")
            if resume_path and isinstance(resume_path, str):
                # Remove leading `/media/` or `/`
                if resume_path.startswith("/media/"):
                    resume_path = resume_path.replace("/media/", "")
                elif resume_path.startswith("/"):
                    resume_path = resume_path[1:]

                # Build full absolute path - try both 'resume' and 'resumes' directories
                possible_paths = [
                    os.path.join(settings.BASE_DIR, 'media/resumes/', os.path.basename(resume_path)),
                    os.path.join(settings.BASE_DIR, 'media/resume/', os.path.basename(resume_path)),
                    os.path.join(settings.MEDIA_ROOT, resume_path) if hasattr(settings, 'MEDIA_ROOT') else None
                ]

                full_path = None
                for path in possible_paths:
                    if path and os.path.exists(path):
                        full_path = path
                        break

                if full_path and os.path.exists(full_path):
                    # Read the file content first
                    with open(full_path, "rb") as f:
                        file_content = f.read()
                    
                    # Create a ContentFile with the file content
                    from django.core.files.base import ContentFile
                    file_obj = ContentFile(file_content, name=os.path.basename(full_path))
                    data["resume"] = file_obj
                else:
                    return Response({
                        "error": "Resume file not found on server",
                        "searched_paths": [p for p in possible_paths if p],
                        "current_dir": os.getcwd()
                    }, status=400)

                if job:
                    try:
                        new_job_id = int(job)
                    except (TypeError, ValueError):
                        new_job_id = None
                else:
                    new_job_id = None

            if serializer.instance.source != '':
                try:
                    try:
                        email_Settings = EmailSettings.objects.get(user_id=job_obj.created_by, company=job_obj.company)
                        if email_Settings:
                            message = f"Thank you for applying to {job_obj.title} at {job_obj.company.name}. We have received your application and will review it shortly."
                            sendMail = CandidateCreateMailer(
                                emails=serializer.instance.webform_candidate_data["Personal Details"]["email"],
                                message=message,
                                email_Settings=email_Settings,
                                candidate=serializer.instance,
                                job=job_obj,
                                account=job_obj.created_by
                            )
                            sendMail.start()
                    except (EmailSettings.DoesNotExist, EmailSettings.MultipleObjectsReturned):
                        # Log that email settings not found for this user
                        pass
                except EmailCategory.DoesNotExist:
                    # Log that email category not found
                    pass
                except Exception as e:
                    # Log any other unexpected errors
                    pass
            
            performed_by = get_safe_performed_by(request)
            if job:
                CandidateTimeline.log_activity(serializer.instance, 'CANDIDATE_CREATED', 'Candidate Created', job=data.get('job'), performed_by=performed_by,pipeline_stage="Screening")
            else:
                CandidateTimeline.log_activity(serializer.instance, 'CANDIDATE_CREATED', 'Candidate Created', performed_by=performed_by)
            if job_obj:    
                JobTimeline.log_activity(job=job_obj,candidate=serializer.instance, activity_type= 'CANDIDATE_JOB_APPLIED', title= 'Candidate Job Applied',description=f"Candidate job applied", performed_by=performed_by)

            
            if assessment_data and job:
                print("....................22")
                # Start assessment processing in background thread
                assessment_thread = threading.Thread(
                    target=run_assessment_background,
                    args=(
                        candidate.id,
                        assessment_data,
                        job.id if isinstance(job, Job) else job,  # Handle both object and ID
                        data.get('resume_data', '')
                    ),
                    daemon=True,
                    name=f"Assessment-{candidate.id}-{getattr(job, 'id', job)}"
                )
                assessment_thread.start()
            else:
                print("....................23")
                rjms_thread = threading.Thread(
                    target=run_rjms_background,
                    args=(candidate.id, job.id, data.get('resume_data', ''), assessment_data,100),
                    daemon=True,
                    name=f"RJMS-Analysis-{candidate.id}-{job.id}"
                )
                rjms_thread.start()
            if credit_wallets:
                for cw in credit_wallets:
                    cw.used_credit += 1
                    cw.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
            
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# class TasksApiView(APIView):
#     def get(self, request, pk=None, *args, **kwargs):
#         if pk is not None:
#             # Retrieve a specific task by ID
#             task = self.get_object(pk)
#             serializer = TasksSerializer(task)
#             return Response(serializer.data)
#         else:
#             # Retrieve all tasks
#             tasks = models.Tasks.objects.all()
#             serializer = TasksSerializer(tasks, many=True)
#             return Response(serializer.data)

#     def post(self, request, *args, **kwargs):
#         # Create a new task
#         serializer = TasksSerializer(data=request.data)
#         if serializer.is_valid():
#             serializer.save()
#             return Response(serializer.data, status=status.HTTP_201_CREATED)
#         return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

#     def put(self, request, pk, *args, **kwargs):
#         # Update a task
#         task = self.get_object(pk)
#         serializer = TasksSerializer(task, data=request.data)
#         if serializer.is_valid():
#             serializer.save()
#             return Response(serializer.data)
#         return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

#     def delete(self, request, pk, *args, **kwargs):
#         # Delete a task
#         task = self.get_object(pk)
#         task.delete()
#         return Response(status=status.HTTP_204_NO_CONTENT)

#     def get_object(self, pk):
#         # Helper method to get a task by pk
#         try:
#             return models.Tasks.objects.get(pk=pk)
#         except models.Tasks.DoesNotExist:
#             raise Http404
 
class TaskOptionsView(APIView):
    def get(self, request, *args, **kwargs):
        # Prepare the options data
        options = {
            "priority_fields": dict(PRIORITY_FIELDS),
            "task_repeat": dict(TASK_REPEAT),
            "task_alert": dict(TASK_ALERT),
            "status_fields": dict(STATUS_FIELDS),
            "currency": dict(CURRENCY),
        }
        # Return the options as JSON response
        return Response(options) 
    
class TaskListCreateView(APIView):
    def get(self, request, id, *args, **kwargs):
        tasks = Tasks.objects.filter(candidate__id=id)
        serializer = TaskSerializer(tasks, many=True)
        return Response(serializer.data)

    def post(self, request, *args, **kwargs):
        serializer = TaskSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            performed_by = get_safe_performed_by(request)
            CandidateTimeline.log_activity(serializer.instance.candidate, 'TASK_CREATED', 'Task Created',description=f"Task created by {request.user.first_name if request.user and hasattr(request.user, 'first_name') else 'Unknown'}", performed_by=performed_by, related_task=serializer.instance)
           
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)  
        
    def put(self, request, id, *args, **kwargs):
        try:
            task = Tasks.objects.get(id=id)
        except Tasks.DoesNotExist:
            return Response({"error": "Task not found"}, status=status.HTTP_404_NOT_FOUND)
        
        serializer = TaskSerializer(task, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def delete(self, request, id, *args, **kwargs):
        try:
            task = Tasks.objects.get(id=id)
        except Tasks.DoesNotExist:
            return Response({"error": "Task not found"}, status=status.HTTP_404_NOT_FOUND)
        
        task.delete()
        return Response({"message": "Task deleted"}, status=status.HTTP_204_NO_CONTENT)

class MarkTaskCompletedView(APIView):
    def post(self, request, id, *args, **kwargs):
        try:
            task = Tasks.objects.get(id=id)
            task.completed = True
            task.status = 'COMPLETED'
            task.save()
            return Response({'status': 'success', 'message': 'Task marked as completed'}, status=status.HTTP_200_OK)
        except Tasks.DoesNotExist:
            return Response({'status': 'error', 'message': 'Task not found'}, status=status.HTTP_404_NOT_FOUND)
        
class EventOptionsView(APIView):
    def get(self, request, *args, **kwargs):
        # Prepare the options data
        options = {
            "meeting_venue": dict(MEETING_VENUE),
            "related_to": dict(RELATED_TO),
            "repeat_fields": dict(REPEAT_FIELDS),
            "reminder_fields": dict(REMINDER_FIELDS),
            "todo_type": dict(TODO_TYPE),
            "currency": dict(CURRENCY),
        }
        # Return the options as JSON response
        return Response(options)
        
class EventListCreateView(APIView):
    def get(self, request, id, *args, **kwargs):
        events = Events.objects.filter(candidate__id=id)
        serializer = EventSerializer(events, many=True)
        return Response(serializer.data)

    def post(self, request, *args, **kwargs):
        serializer = EventSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            performed_by = get_safe_performed_by(request)
            CandidateTimeline.log_activity(serializer.instance.candidate, 'EVENT_CREATED', 'Event Created',description=f"Event created by {request.user.account_id if request.user and hasattr(request.user, 'account_id') else 'Unknown'}", performed_by=performed_by, related_event=serializer.instance)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def put(self, request, id, *args, **kwargs):
        try:
            event = Events.objects.get(id=id)
        except Events.DoesNotExist:
            return Response({"error": "Event not found"}, status=status.HTTP_404_NOT_FOUND)

        serializer = EventSerializer(event, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, id, *args, **kwargs):
        try:
            event = Events.objects.get(id=id)
        except Events.DoesNotExist:
            return Response({"error": "Event not found"}, status=status.HTTP_404_NOT_FOUND)

        event.delete()
        return Response({"message": "Event deleted"}, status=status.HTTP_204_NO_CONTENT)

class CallOptionsView(APIView):
    def get(self, request, *args, **kwargs):
        # Prepare the options data
        options = {
            "call_type": dict(CALL_TYPE_CHOICES),
            "call_purpose": dict(CALL_PURPOSE_CHOICES),
            "call_status": dict(CALL_STATUS),
            "reminder_fields": dict(CALL_REMINDER_FIELDS),
            "todo_type": dict(TODO_TYPE),
        }
        # Return the options as JSON response
        return Response(options)

class CallListCreateView(APIView):
    def get(self, request, id, *args, **kwargs):
        calls = Call.objects.filter(candidate__id=id)
        serializer = CallSerializer(calls, many=True)
        return Response(serializer.data)

    def post(self, request, *args, **kwargs):
        owner_id = request.data.get('owner')
        if owner_id:
            try:
                owner_instance = Account.objects.get(account_id=owner_id)
                request.data['owner'] = owner_instance.pk  # Assign only the primary key
            except Account.DoesNotExist:
                return Response({'error': 'Owner not found'}, status=status.HTTP_400_BAD_REQUEST)
        
        serializer = CallSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            performed_by = get_safe_performed_by(request)
            CandidateTimeline.log_activity(serializer.instance.candidate, 'CALL_CREATED', 'Call Created',description=f"Call created", performed_by=performed_by, related_call=serializer.instance)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def put(self, request, id, *args, **kwargs):
        try:
            call = Call.objects.get(id=id)
        except Call.DoesNotExist:
            return Response({"error": "Call not found"}, status=status.HTTP_404_NOT_FOUND)
        
        owner_id = request.data.get('owner')
        if owner_id:
            try:
                owner_instance = Account.objects.get(account_id=owner_id)
                request.data['owner'] = owner_instance.pk  # Assign only the primary key
            except Account.DoesNotExist:
                return Response({'error': 'Owner not found'}, status=status.HTTP_400_BAD_REQUEST)

        serializer = CallSerializer(call, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, id, *args, **kwargs):
        try:
            call = Call.objects.get(id=id)
        except Call.DoesNotExist:
            return Response({"error": "Call not found"}, status=status.HTTP_404_NOT_FOUND)

        call.delete()
        return Response({"message": "Call deleted"}, status=status.HTTP_204_NO_CONTENT)


# class TaskByCandidate(APIView):
#     def get(self, request, pk=None, *args, **kwargs):
#         if pk is not None:
#             # Use get_object_or_404 to retrieve the candidate or return 404 if not found
#             candidate = get_object_or_404(models.Candidate, pk=pk)            
#             tasks = models.Tasks.objects.filter(candidate=candidate)
#             serializer = TasksSerializer(tasks, many=True)
#             return Response(serializer.data)

#         else:
#             # If no primary key is provided, return a "not found" response
#             return Response({"detail": "Candidate not found"}, status=404)





# resume_parser/views.py
import base64
import json
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
import base64
import requests
import json

class ParseResumeAPIView(APIView):
    def post(self, request, *args, **kwargs):
        user_key = 'NFSF6Y36'
        version = '8.0.0'
        sub_user_id = 'NFSF6Y36'
        APIURL="https://rest.rchilli.com/RChilliParser/Rchilli/parseResumeBinary"
        file_data = request.data.get('resume')
        file_name = file_data.name  # Retrieve the filename from the file object

        encoded_string = base64.b64encode(file_data.read())
        data64 = encoded_string.decode('UTF-8')

        headers = {'content-type': 'application/json'}

        body = {
            "filedata": data64,
            "filename": file_name,
            "userkey": user_key,
            "version": version,
            "subuserid": sub_user_id
        }

        try:
            response = requests.post(APIURL, data=json.dumps(body), headers=headers)
            response.raise_for_status()  # Raise an HTTPError for bad responses
            resp_data = response.json().get("ResumeParserData", {})

            return Response(resp_data, status=status.HTTP_200_OK)

        except requests.exceptions.RequestException as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
  
def calculate_cost(model, usage):
    """
    Calculate estimated OpenAI API cost based on token usage.

    model: the model name you passed to the API
    usage: the .usage object returned by OpenAI
           (with .prompt_tokens, .completion_tokens, and possibly .training_tokens)

    Returns a dict:
      {
        "model":            resolved_model_key,
        "input_tokens":     int,
        "output_tokens":    int,
        "training_tokens":  int,      # always 0 for chat completions
        "input_cost":       float,    # $ rounded to 6 decimals
        "output_cost":      float,    # $ rounded to 6 decimals
        "training_cost":    float,    # $ rounded to 6 decimals
        "total_cost":       float,    # $ rounded to 6 decimals
      }
    or None if the model is unrecognized.
    """
    # Pricing in $ per **1 000 000** tokens (fine-tuning prices from OpenAI UI)
    pricing_per_million = {
        "gpt-4o-mini-2024-07-18": {
            "input":    0.30,   # $0.30 / 1M tokens
            "cached":   0.15,   # $0.15 / 1M tokens  (if you ever track cached vs new)
            "output":   1.20,   # $1.20 / 1M tokens
            "training": 3.00    # $3.00  / 1M tokens (for fine-tune runs)
        },
        "gpt-3.5-turbo-1106": {
            "input":    0.0010 * 1000,  # convert from $0.0010 / 1K → $1.0 / 1M
            "output":   0.0020 * 1000,  # convert from $0.0020 / 1K → $2.0 / 1M
            "training": 0.0
        },
    }

    # resolve exact key (allow variants like "gpt-4o-mini-2024-07-18-xyz")
    resolved = None
    for key in pricing_per_million:
        if model.startswith(key):
            resolved = key
            break
    if not resolved:
        return None

    rates = pricing_per_million[resolved]
    # convert $ / 1 000 000 tokens → $ / 1 token
    input_rate    = rates["input"]    / 1_000_000
    output_rate   = rates["output"]   / 1_000_000
    training_rate = rates.get("training", 0.0) / 1_000_000

    in_tokens    = getattr(usage, "prompt_tokens",     0)
    out_tokens   = getattr(usage, "completion_tokens", 0)
    train_tokens = getattr(usage, "training_tokens",   0)

    in_cost    = in_tokens    * input_rate
    out_cost   = out_tokens   * output_rate
    train_cost = train_tokens * training_rate
    total_cost = in_cost + out_cost + train_cost

    return {
        "model":           resolved,
        "input_tokens":    in_tokens,
        "output_tokens":   out_tokens,
        "training_tokens": train_tokens,
        "input_cost":      round(in_cost,    6),
        "output_cost":     round(out_cost,   6),
        "training_cost":   round(train_cost, 6),
        "total_cost":      round(total_cost, 6),
    }
   
         
class ExtractResumeView(APIView):
    def post(self, request, *args, **kwargs):
        try:
            pdf_file = request.data.get('resume')
            if not pdf_file:
                return Response({'error': 'Resume file is required'}, status=status.HTTP_400_BAD_REQUEST)

            # Get file content and determine file type
            file_content = pdf_file.read()
            file_name = pdf_file.name.lower()
            
            # Extract text based on file type
            text_content = self.extract_text_from_file(file_content, file_name)
            
            if not text_content:
                return Response({'error': 'Could not extract text from file'}, status=status.HTTP_400_BAD_REQUEST)

            # 1) load your JSON-template
            with open('resume_json_format.json') as f:  
                template_dict = json.load(f)
				
			# 3) Define every repeatable‐array field you want to be dynamic
            dynamic_keys = [
				'Email',
				'PhoneNumber',
				'Address',
				'SegregatedQualification',
				'SegregatedSkill',
				'SegregatedExperience',
			]

			# 4) For each, keep exactly one "shape" element (or [] if none)
            for key in dynamic_keys:
                arr = template_dict.get(key, [])
                template_dict[key] = arr[:1] if isinstance(arr, list) else []

            full_json_template = json.dumps(template_dict, indent=2)

			# Build prompt: return EVERY field from schema, no omissions
            prompt = (
				"You are an expert resume parser.\n"
				"Return a single JSON object that matches the schema EXACTLY.\n"
				"Rules:\n"
				"- Include EVERY key at every nesting level from the schema; do not drop or rename keys.\n"
				"- If a value is missing in the resume, use defaults based on the schema:\n"
				"  - Strings: \"\"\n"
				"  - Arrays: [] (include 0 items)\n"
				"  - Objects: include all child keys with \"\" (and [] where applicable).\n"
				"- Do not invent data; only extract what the resume states.\n"
				"- For array fields, repeat the included template element for each found item; if none, return [].\n"
				"- Output raw JSON only; no markdown or commentary.\n\n"
				"Additionally, populate these professional fields at the root of the JSON:\n"
				"- YearsOfExperience: total years across roles (decimal allowed, e.g., \"3.5\").\n"
				"- ProfessionalStartDate: earliest StartDate among experiences.\n"
				"- ProfessionalEndDate: EndDate of the most recent role; leave \"\" if currently working.\n"
				"- CurrentlyWorking: \"Yes\" if any role is ongoing/marked current; else \"No\".\n"
				"- CurrentJobTitle: job title of the most recent role.\n"
				"- HighestQualification: highest degree attained.\n"
				"- ProfessionalDegree: the degree name of the highest qualification (or concise summary).\n"
				"- FunctionalArea: infer from roles/skills when stated; else \"\".\n"
				"- NoticePeriod: extract if present; else \"\".\n"
				"- SalaryCurrency: currency code/symbol mentioned (or from CurrentSalary.Currency if present).\n"
				"- ProfessionalCertificate: key certifications; else \"\".\n"
				"- Subject: main specialization/subject; else \"\".\n"
				"- CurriculumBoard: education board if stated; else \"\".\n\n"
				"For phone numbers, format them as follows:\n"
				"- Add country code (+91) for Indian numbers\n"
				"- Format as: +91-XXXXXXXXXX\n"
				"- Set FormattedNumber to the formatted version\n\n"
				"Resume Text:\n"
				f"{text_content}\n\n"
				"JSON Schema:\n"
				f"{full_json_template}"
			)

            completion = client.chat.completions.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[
                    {"role": "system", "content": "You are an AI resume parser."},
                    {"role": "user",   "content": prompt}
                ]
            )

            # 1) grab and clean the raw text
            response_content = completion.choices[0].message.content.strip()
            if response_content.startswith("```"):
                # strip ``` … ```
                response_content = response_content.split("```")[-2].strip()

            # 3) attempt to parse as JSON
            try:
                parsed_data = json.loads(response_content)
            except json.JSONDecodeError as json_error:
                # return the raw payload alongside the parse error
                return Response({
                    "error":           "Invalid JSON from AI",
                    "details":         str(json_error),
                    "raw_response":    response_content
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            return Response(parsed_data, status=status.HTTP_200_OK)

        except Exception as e:
            return Response(
                {"error": f"Error processing file: {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST
            )
    
    def extract_text_from_file(self, file_content, file_name):
        """Extract text from PDF and DOCX files only"""
        try:
            # Determine file type
            if file_name.endswith('.pdf'):
                return self.extract_text_from_pdf(file_content)
            elif file_name.endswith(('.docx', '.doc')):
                return self.extract_text_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_name}. Only PDF and DOCX files are supported.")
        except Exception as e:
            logger.warning(f"Error extracting text: {str(e)}")
            return None
    
    def extract_text_from_pdf(self, file_content):
        """Extract text from PDF file"""
        try:
            return extract_text(io.BytesIO(file_content))
        except Exception as e:
            logger.warning(f"Error extracting PDF text: {str(e)}")
            return None
    
    def extract_text_from_docx(self, file_content):
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n'.join(text_content)
        except Exception as e:
            logger.warning(f"Error extracting DOCX text: {str(e)}")
            return None

from rest_framework.parsers import JSONParser, MultiPartParser
# Live progress (SSE) for resume extraction
class ExtractResumeProgressSSE(APIView):
    # @handle_ai_credits(feature_code="AI_CREDITS", usage_code="Resume_Parcing_(using AI)")
    def post(self, request, *args, **kwargs):
        try:
            pdf_file = request.data.get('resume')

            if not pdf_file:
                return Response({'error': 'Resume file is required'}, status=status.HTTP_400_BAD_REQUEST)

            def event_stream():
                try:
                    # 10%: file accepted
                    yield f"event: progress\ndata: {json.dumps({'percent': 10, 'message': 'File received'})}\n\n"

                    file_content = pdf_file.read()
                    file_name = pdf_file.name.lower()

                    # 30%: extracting text
                    yield f"event: progress\ndata: {json.dumps({'percent': 30, 'message': 'Extracting text'})}\n\n"
                    text_content = ExtractResumeView().extract_text_from_file(file_content, file_name)
                    if not text_content:
                        yield f"event: error\ndata: {json.dumps({'message': 'Could not extract text from file'})}\n\n"
                        return

                    # 40%: loading schema
                    yield f"event: progress\ndata: {json.dumps({'percent': 40, 'message': 'Loading JSON schema'})}\n\n"
                    with open('resume_json_format.json') as f:
                        template_dict = json.load(f)

                    dynamic_keys = [
                        'Email',
                        'PhoneNumber',
                        'Address',
                        'SegregatedQualification',
                        'SegregatedSkill',
                        'SegregatedExperience',
                    ]
                    for key in dynamic_keys:
                        arr = template_dict.get(key, [])
                        template_dict[key] = arr[:1] if isinstance(arr, list) else []

                    # 50%: building prompt
                    yield f"event: progress\ndata: {json.dumps({'percent': 50, 'message': 'Building prompt'})}\n\n"
                    full_json_template = json.dumps(template_dict, indent=2)
                    prompt = (
                        "You are an expert resume parser.\n"
                        "Return a single JSON object that matches the schema EXACTLY.\n"
                        "Rules:\n"
                        "- Include EVERY key at every nesting level from the schema; do not drop or rename keys.\n"
                        "- If a value is missing in the resume, use defaults based on the schema:\n"
                        "  - Strings: \"\"\n"
                        "  - Arrays: [] (include 0 items)\n"
                        "  - Objects: include all child keys with \"\" (and [] where applicable).\n"
                        "- Do not invent data; only extract what the resume states.\n"
                        "- For array fields, repeat the included template element for each found item; if none, return [].\n"
                        "- Output raw JSON only; no markdown or commentary.\n\n"
                        "Additionally, populate these professional fields at the root of the JSON:\n"
                        "- YearsOfExperience: total years across roles (decimal allowed, e.g., \"3.5\").\n"
                        "- ProfessionalStartDate: earliest StartDate among experiences.\n"
                        "- ProfessionalEndDate: EndDate of the most recent role; leave \"\" if currently working.\n"
                        "- CurrentlyWorking: \"Yes\" if any role is ongoing/marked current; else \"No\".\n"
                        "- CurrentJobTitle: job title of the most recent role.\n"
                        "- HighestQualification: highest degree attained.\n"
                        "- ProfessionalDegree: the degree name of the highest qualification (or concise summary).\n"
                        "- FunctionalArea: infer from roles/skills when stated; else \"\".\n"
                        "- NoticePeriod: extract if present; else \"\".\n"
                        "- SalaryCurrency: currency code/symbol mentioned (or from CurrentSalary.Currency if present).\n"
                        "- ProfessionalCertificate: key certifications; else \"\".\n"
                        "- Subject: main specialization/subject; else \"\".\n"
                        "- CurriculumBoard: education board if stated; else \"\".\n\n"
                        "For phone numbers, format them as follows:\n"
                        "- Add country code (+91) for Indian numbers\n"
                        "- Format as: +91-XXXXXXXXXX\n"
                        "- Set FormattedNumber to the formatted version\n\n"
                        f"Resume Text:\n{text_content}\n\n"
                        f"JSON Schema:\n{full_json_template}"
                    )

                    # 70%: calling model
                    yield f"event: progress\ndata: {json.dumps({'percent': 70, 'message': 'Calling AI model'})}\n\n"
                    # combined_prompt = "\n\n".join([m["content"] for m in messages])
                    API_KEY = "AIzaSyAuIJJFzcq4RYX69G4_XdwFYcVWri-QNMY"
                    # client_gemini = genai.Client(api_key=API_KEY)
                    # response_gemini = client_gemini.models.generate_content(
                    #     model="gemini-2.5-flash-lite",
                    #     contents=prompt
                    # )

                    completion = client.chat.completions.create(
                        model="gpt-4o-mini-2024-07-18",
                        messages=[
                            {"role": "system", "content": "You are an AI resume parser."},
                            {"role": "user",   "content": prompt}
                        ]
                    )
                    input_tokens = completion.usage.prompt_tokens
                    output_tokens = completion.usage.completion_tokens

                    # 🔹 Cost (update if pricing changes)
                    input_cost = (input_tokens / 1_000_000) * 5.00
                    output_cost = (output_tokens / 1_000_000) * 15.00

                    print(
                        f"Resume Extract → in:{input_tokens} out:{output_tokens} "
                        f"cost:${input_cost + output_cost:.6f}"
                    )


                    # 80%: got response
                    yield f"event: progress\ndata: {json.dumps({'percent': 80, 'message': 'Processing AI response'})}\n\n"
                    response_content = completion.choices[0].message.content.strip()
                    if response_content.startswith("```"):
                        response_content = response_content.split("```")[-2].strip()

                    try:
                        parsed_data = json.loads(response_content)
                        # Handle Gemini response - might be wrapped in markdown code blocks
                        # gemini_data = None
                        # if response_gemini.text and response_gemini.text.strip():
                        #     try:
                        #         gemini_text = response_gemini.text.strip()
                        #         # Check if response is wrapped in markdown code blocks
                        #         if gemini_text.startswith("```json") and gemini_text.endswith("```"):
                        #             # Extract JSON from between the code blocks
                        #             json_start = gemini_text.find("\n") + 1
                        #             json_end = gemini_text.rfind("```")
                        #             gemini_text = gemini_text[json_start:json_end].strip()

                        #         gemini_data = json.loads(gemini_text)
                        #     except json.JSONDecodeError:
                        #         # If still not valid JSON, store as text
                        #         gemini_data = {"raw_response": response_gemini.text}
                        # else:
                        #     gemini_data = {"raw_response": "No response from Gemini"}
                    except json.JSONDecodeError as json_error:
                        yield f"event: error\ndata: {json.dumps({'message': 'Invalid JSON from AI', 'details': str(json_error)})}\n\n"
                        return

                    # 100%: done
                    yield f"event: progress\ndata: {json.dumps({'percent': 100, 'message': 'Completed'})}\n\n"

                    yield f"event: result\ndata: {json.dumps(parsed_data)}\n\n"
                    # yield f"event: result\ndata: {json.dumps(parsed_data)}\n\n{json.dumps({'parsed_data': parsed_data, 'usage': f'Job Description -> in:{input_tokens} out:{output_tokens} cost:${input_cost + output_cost:.6f}', 'gemini_data': gemini_data, 'gemini_usage': serialize_usage_metadata(response_gemini.usage_metadata)})}\n\n"
                except Exception as ex:
                    yield f"event: error\ndata: {json.dumps({'message': str(ex)})}\n\n"

            response = StreamingHttpResponse(event_stream(), content_type='text/event-stream')
            response['Cache-Control'] = 'no-cache'
            response['X-Accel-Buffering'] = 'no'  # for proxies like nginx
            return response
        except Exception as e:
            return Response(
                {"error": f"Error processing file: {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST
            )
            
# class CreateMultipleCandidates(APIView):
#     parser_classes = [JSONParser, MultiPartParser]

#     def post(self, request, candidate_id, *args, **kwargs):
#         # Get the existing candidate by ID
#         existing_candidate = models.Candidate.objects.filter(id=candidate_id).first()

#         if not existing_candidate:
#             return Response({'error': 'Candidate not found'}, status=status.HTTP_404_NOT_FOUND)

#         # Get the data of the existing candidate
#         existing_candidate_data = MultipleCandidateSerializer(existing_candidate).data

#         # Assuming the request data contains a list of job IDs
#         job_ids = request.data.get('job_ids', [])

#         created_candidates = []

#         for job_id in job_ids:
#             # Check if a candidate already exists for the current job ID with the same first name
#             existing_candidate_for_job = models.Candidate.objects.filter(
#                 job_id=job_id,
#                 first_name=existing_candidate.first_name
#             ).first()

#             if existing_candidate_for_job:
#                 # Skip creating a new candidate if one already exists for the job with the same first name
#                 created_candidates.append(existing_candidate_for_job)
#             else:
#                 # Create a new candidate object for each job
#                 new_candidate_data = existing_candidate_data.copy()
#                 new_candidate_data['job'] = job_id  # Assign the current job ID

#                 serializer = MultipleCandidateSerializer(data=new_candidate_data)
#                 if serializer.is_valid():
#                     # Save the new candidate
#                     candidate = serializer.save()

#                     # Get the newly created candidate object
#                     new_candidate = models.Candidate.objects.get(pk=candidate.id)

#                     # Copy relevant fields from existing candidate to new candidate
#                     new_candidate.resume = existing_candidate.resume
#                     new_candidate.cover_letter = existing_candidate.cover_letter
#                     new_candidate.certificate = existing_candidate.certificate
#                     new_candidate.professional_certificate = existing_candidate.professional_certificate
#                     new_candidate.save()

#                     created_candidates.append(new_candidate)
#                 else:
#                     print(serializer.errors)

#         return Response(status=status.HTTP_201_CREATED)

class CreateMultipleCandidates(APIView):
    parser_classes = [JSONParser, MultiPartParser]

    def post(self, request, candidate_id, *args, **kwargs):
        # Get the existing candidate by ID
        existing_candidate = Candidate.objects.filter(id=candidate_id).first()

        if not existing_candidate:
            return Response({'error': 'Candidate not found'}, status=status.HTTP_404_NOT_FOUND)

        # Assuming the request data contains a list of job IDs
        job_ids = request.data.get('job_ids', [])

        for job_id in job_ids:
            job = Job.objects.get(id=job_id)
            existing_candidate.job.add(job)

        existing_candidate.save()
        performed_by = get_safe_performed_by(request)
        CandidateTimeline.log_activity(existing_candidate, 'CANDIDATE_JOB_ASSIGNED', 'Candidate Job Assigned', job=job_ids, performed_by=performed_by)
        JobTimeline.log_activity(job=job,candidate=existing_candidate, activity_type= 'CANDIDATE_JOB_ASSIGNED', title= 'Candidate Job Assigned',description=f"Candidate job assigned", performed_by=performed_by)

        return Response(status=status.HTTP_200_OK)
    
class CandidateUpload_documentApi(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        data = helper.uploadCandidateDocument(request)
        return Response(data, status=data.get("status", status.HTTP_200_OK))

    def delete(self, request):
        data = helper.deleteCandidateDocument(request)
        return Response(data, status=data.get("status", status.HTTP_200_OK))
    

class EmailSettingsApiView(APIView):
    
    def get(self, request, *args, **kwargs):
        
        user = request.user

        if user.is_superuser:
            email_settings = EmailSettings.objects.all()
        else:
            email_settings = EmailSettings.objects.filter(user_id=user.id)
        
        serializer = EmailSettingsSerializer(email_settings, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
        
    @check_feature_availability('Email_SMTP_IMPS')
    def post(self, request, *args, **kwargs):
        
        serializer = EmailSettingsSerializer(data=request.data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def put(self, request, *args, **kwargs):
        response = helper.updateEmailSmtpSetting(request)
        return response
    
    def delete(self, request, *args, **kwargs):
        email_setting_id = kwargs.get('pk')
        response = helper.deleteEmailSmtpSetting(request, email_setting_id)
        return response

class DailyCandidateCountApi(APIView):
    """
    Get daily created candidate counts and statistics for a given date range.
    
    Query params:
    - days: Number of days to look back (default: 7)
    """
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def get(self, request):
        """
        Handle GET request to retrieve daily candidate statistics.
        """
        from . import helper
        data = helper.get_daily_candidate_counts(request)
        return Response(data)

class ExecuteSQLQueryView(APIView):
    def post(self, request, *args, **kwargs):
        try:
            user = request.user
            data = request.data
            query = data.get('query')
            if query:
                query = query.replace('%', '')

            # Use Django ORM to filter candidates
            candidates = Candidate.objects.select_related('country', 'state', 'city',  'professional_degree', 'current_job_title', 'webform').prefetch_related('subjects', 'job')
            
            if not user.is_superuser:
                candidates = candidates.filter(job__company__id=user.company_id)

            if query:
                # Replace logical operators with Python's operators
                query = query.replace(" AND ", " and ").replace(" OR ", " or ")
                
                # Parse and construct the Q object
                q_object = self.construct_q_object(query)
                
                # Apply the Q object to filter candidates
                candidates = candidates.filter(q_object)

            results = candidates.values(
                'id', 'first_name', 'last_name', 'middle_name', 'pipeline_stage_status', 'gender', 'marital_status', 'email', 'mobile', 'source', 'notice_period', 'exp_years', 'highest_qualification', 'professional_start_date', 'professional_end_date', 'current_job_title__name', 'professional_degree__name','job__title', 'fun_area', 'curriculum_board', 'updated',
                'country__name', 'state__name', 'city__name'
            )

            results_list = list(results)
            
            for candidate in candidates:
                candidate_data = next((item for item in results_list if item['id'] == candidate.id), None)
                if candidate_data:
                    candidate_data['subjects'] = [subject.name for subject in candidate.subjects.all()] 

            return Response({'result': results_list}, status=status.HTTP_200_OK)
        except Exception as e:
            # Return all candidates in case of an error
            candidates = Candidate.objects.select_related(
                'country', 'state', 'city', 'professional_degree', 'current_job_title', 'webform'
            ).prefetch_related('subjects', 'job').values(
                'id', 'first_name', 'last_name', 'middle_name', 'pipeline_stage_status', 'gender', 'marital_status', 'email', 'mobile', 'source', 'notice_period', 'exp_years', 'highest_qualification', 'professional_start_date', 'professional_end_date', 'job__title', 'current_job_title__name', 'professional_degree__name', 'fun_area', 'curriculum_board', 'updated',
                'country__name', 'state__name', 'city__name'
            )
            results_list = list(candidates)
            return Response({'result': results_list, 'error': str(e)}, status=status.HTTP_200_OK)

    def construct_q_object(self, query):
        # Tokenize the query using regex to handle nested conditions and operators
        tokens = re.split(r'(\(|\)|\s+and\s+|\s+or\s+)', query)
        tokens = [token for token in tokens if token.strip()]

        # Stack to manage nested conditions and Q objects
        stack = []
        current_q = None
        current_operator = None

        for token in tokens:
            token = token.strip().lower()

            if token == '(':
                stack.append((current_q, current_operator))
                current_q = None
                current_operator = None
            elif token == ')':
                temp_q = current_q
                if stack:
                    current_q, current_operator = stack.pop()
                    if current_operator == 'and':
                        current_q &= temp_q
                    elif current_operator == 'or':
                        current_q |= temp_q
                    else:
                        current_q = temp_q
            elif token == 'and':
                current_operator = 'and'
            elif token == 'or':
                current_operator = 'or'
            else:
                # Parse the condition token (e.g., "field operator value")
                # match = re.match(r'(\w+) (like|not like|=|<>|>|>=|<|<=|startwith|endwith) (.+)', token)
                match = re.match(r"(\w+)\s+(like|not like|=|<>|>|>=|<|<=|startwith|endwith|IS NULL|IS NOT NULL)\s+('(?:[^'\\]|\\.)*'|\"(?:[^\"\\]|\\.)*\"|[^\s]+)", token, re.IGNORECASE)
                if match:
                    field, operator, value = match.groups()
                    value = value.strip().strip("'")
                    
                    # Map field names to Django ORM fields
                    if field in ['country', 'state', 'city']:
                        field = f"{field}__name"
                    elif field == 'job':
                        field = 'job__title'
                    
                    # Construct the condition Q object
                    if operator == 'like':
                        condition = Q(**{f"{field}__icontains": value})
                    elif operator == 'not like':
                        condition = ~Q(**{f"{field}__icontains": value})
                    elif operator == '=':
                        condition = Q(**{f"{field}__iexact": value})
                    elif operator == '<>':
                        condition = ~Q(**{f"{field}__iexact": value})
                    elif operator == '>':
                        condition = Q(**{f"{field}__gt": value})
                    elif operator == '>=':
                        condition = Q(**{f"{field}__gte": value})
                    elif operator == '<':
                        condition = Q(**{f"{field}__lt": value})
                    elif operator == '<=':
                        condition = Q(**{f"{field}__lte": value})
                    elif operator == 'startwith':
                        condition = Q(**{f"{field}__istartswith": value})
                    elif operator == 'endwith':
                        condition = Q(**{f"{field}__iendswith": value})
                    elif operator == 'is null':
                        condition = Q(**{f"{field}__isnull": True})
                    elif operator == 'is not null':
                        condition = ~Q(**{f"{field}__isnull": True})

                    # Apply the current operator
                    if current_q is None:
                        current_q = condition
                    elif current_operator == 'and':
                        current_q &= condition
                    elif current_operator == 'or':
                        current_q |= condition

                    current_operator = None

        return current_q if current_q else Q()

def AddData(request):
    # Load the Excel file
    excel_file_path = os.path.join(settings.BASE_DIR, 'degree_names_list.xlsx')
    
    sheet_model_mapping = {
        # 'departments': {
        #     'model': Department,
        #     'fields': {'name': 'Departments'}
        # },
        # 'Degree': {
        #     'model': Degree,
        #     'fields': {'name': 'Degree Name'}
        # },
        'Job Title': {
            'model': Designation,
            'fields': {'name': 'Job Title'}
        },
        # 'Subject Specialization': {
        #     'model': SubjectSpecialization,
        #     'fields': {'name': 'Subject or Specialization'}
        # },
        'skills': {
            'model': Skill,
            'fields': {'name': 'Core Pedagogical Skills:'}
        }
    }
    
    for sheet_name, model_info in sheet_model_mapping.items():
        df = pd.read_excel(excel_file_path, sheet_name=sheet_name)
        
        for index, row in df.iterrows():
            model_instance = model_info['model'](
                name=row[model_info['fields']['name']]
            )
            model_instance.save()
    
    return JsonResponse({'status': 'success', 'message': 'Data added successfully'})


class CandidateWithoutLoginAPI(APIView):
    def post(self, request, *args, **kwargs):
        # Create a new candidate
        data = request.data.copy()
        job_ids = data.pop('job', None)
        
        if not job_ids:
            return Response({"error": "Job is required"}, status=status.HTTP_400_BAD_REQUEST)
        
        serializer = CandidateWithoutLoginSerializer(data=data)
        if serializer.is_valid():
            candidate = serializer.save()
            
            if job_ids is not None:
                jobs = Job.objects.filter(id__in=job_ids)  # Filter jobs by the received IDs
                candidate.job.set(jobs)
                 
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class SkillApiView(APIView):
    def get(self, request, *args, **kwargs):
        data = helper.getSkills(request)
        return makeResponse(data)   

class SubjectSpecializationApiView(APIView):
    def get(self, request, *args, **kwargs):
        data = helper.getSubjectSpecialization(request)
        return makeResponse(data)   
    
# class CandidateApplyJobCareerApi(APIView):
#     def get(self, request,login_email, *args, **kwargs):
#         data = helper.getCandidateApplyJobCareer(request,login_email)
#         return makeResponse(data)
    
class TimeLineView(APIView):
    def get(self, request,candidate_id):
        data = helper.getTimeLine(request,candidate_id)
        return makeResponse(data)

class AssignJobToCandidates(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        data = helper.assignJobToCandidates(request)
        return makeResponse(data)
    
class CandidateBulkMailApi(APIView):
    def post(self, request, *args, **kwargs):
        data = helper.getCandidateBulkMail(request)
        return makeResponse(data)

class DailyEmailQuotaApi(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def get(self, request, *args, **kwargs):
        data = helper.getDailyEmailQuota(request)
        return makeResponse(data)


class HiringStatusTimeLineApi(APIView):
    def get(self, request, id,*args, **kwargs):
        data = helper.HiringStatusTimeLine(request,id)
        return data

class CandidateApplyJobCareerApi(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def get(self, request, *args, **kwargs):
        data = helper.CandidateApplyJobCareer(request)
        return Response(data)

class SaveJobView(APIView):
    permission_classes = [IsAuthenticated]
    authentication_classes = [JWTAuthentication]

    def get(self, request):
        response_data, status_code = helper.getSavedJob(request)
        return Response(response_data, status=status_code)

    def post(self, request):
        response_data, status_code = helper.SaveJob(request)
        return Response(response_data, status=status_code)
    
    def delete(self, request,pk):
        response_data, status_code = helper.DeleteSavedJob(request,pk)
        return Response(response_data, status=status_code)

class CandidateResumeView(APIView):
    def get(self, request):
        response_data, status_code = helper.getCandidateResume(request)
        return Response(response_data, status=status_code)
    
    def post(self, request):
        response_data, status_code = helper.CreateCandidateResume(request)
        return Response(response_data, status=status_code)
    
    def put(self, request,pk):
        response_data, status_code = helper.UpdateCandidateResume(request,pk)
        return Response(response_data, status=status_code)
    
    def delete(self, request,pk):
        response_data, status_code = helper.DeleteCandidateResume(request,pk)
        return Response(response_data, status=status_code)


class CandidateProfileView(APIView):
    def get(self, request):
        response_data, status_code = helper.getCandidateProfile(request)
        return Response(response_data, status=status_code)
    
    def post(self, request):
        response_data, status_code = helper.CreateCandidateProfile(request)
        return Response(response_data, status=status_code)
    
    def put(self, request,pk):
        response_data, status_code = helper.UpdateCandidateProfile(request,pk)
        return Response(response_data, status=status_code)
    
    def delete(self, request,pk):
        response_data, status_code = helper.DeleteCandidateProfile(request,pk)
        return Response(response_data, status=status_code)

class Test(APIView):
    def get(self, request):
        # Get the module form structure
        module = get_object_or_404(Module, id=79)
        
        # Sample resume data (in a real scenario, this would come from parsing a resume)
        resume_data = {
            "Name": {
                "FullName": "Hardeep Singh",
                "LastName": "Singh",
                "FirstName": "Hardeep",
                "MiddleName": ""
            },
            "Email": [{"EmailAddress": "hardeepspm@gmail.com"}],
            "PhoneNumber": [{"FormattedNumber": "+91-802207202"}],
            "DateOfBirth": "1985-05-15",
            "Gender": "Male",
            "MaritalStatus": "Married",
            "YearsOfExperience": "8.0",
            "HighestQualification": "MBA",
            "CurrentlyWorking": "Yes",
            "CurrentJobTitle": "Senior Product Manager",
            "ProfessionalDegree": "MBA - Product Management",
            "ProfessionalStartDate": "2017-03-01",
            "ProfessionalEndDate": "",
            "Address": [{
                "Street": "123 Main St",
                "City": "Bangalore",
                "State": "Karnataka",
                "Country": "India",
                "ZipCode": "560001"
            }],
            "CurrentEmployer": "ZenAdmin.ai",
            "FunctionalArea": "Product Management",
            "NoticePeriod": "2 months",
            "CurrentSalary": {"Amount": "3000000", "Currency": "INR"},
            "ExpectedSalary": {"Amount": "3500000", "Currency": "INR"},
            "SalaryCurrency": "INR",
            "ProfessionalCertificate": "PMP, CSPO",
            "Subject": "Product Management, Agile, Scrum",
            "SegregatedSkill": [
                {"Skill": "Product Development Lifecycle"},
                {"Skill": "Roadmapping & GTM Strategy"},
                {"Skill": "KPIs building and monitoring"}
            ],
            "SegregatedExperience": [
                {
                    "Employer": {"EmployerName": "ZenAdmin.ai"},
                    "JobProfile": {"FormattedName": "Senior Product Manager"},
                    "JobDescription": "Owned end-to-end product for AI-driven Performance Management and HRIS.",
                    "StartDate": "2024-01-01",
                    "EndDate": "2025-06-30"
                },
                {
                    "Employer": {"EmployerName": "Amala.Earth"},
                    "JobProfile": {"FormattedName": "Product Lead"},
                    "JobDescription": "Implemented new 3PL integration, reducing delivery time from 3 days to <24 hours.",
                    "StartDate": "2022-02-01",
                    "EndDate": "2023-05-31"
                }
            ],
            "SegregatedQualification": [
                {
                    "Degree": {"DegreeName": "MBA"},
                    "Institution": {"Name": "Great Lakes Institute of Management"},
                    "StartDate": "2017-01-01",
                    "EndDate": "2017-12-31"
                },
                {
                    "Degree": {"DegreeName": "Bachelor of Commerce"},
                    "Institution": {"Name": "University of Delhi"},
                    "StartDate": "2004-01-01",
                    "EndDate": "2007-12-31"
                }
            ]
        }
        
        # Map resume data to module format
        mapped_data = map_resume_to_module(resume_data, module.form)
        source = "Mailbox"
        job= "124"
        pipeline_stage_status=Associated
        pipeline_stage= Applied
        resume= ""
        webform_candidate_data= mapped_data
        
        return Response({
            'status': 'success',
            'data': mapped_data
        })


class CandidateStatusBySiteAPI(APIView):
    """
    API to get candidate statuses filtered by site and current user
    """
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def get(self, request, *args, **kwargs):
        # Get the current user
        user = request.user
        
        # Get site_id from query params
        try:
            # Get candidates associated with the current user and filter by Interview stage
            candidates = Candidate.objects.filter(
                account=user,
                pipeline_stage='Interview'  # Filter by Interview stage
            )
            # Get interviews for the filtered candidates
            interviews = Interview.objects.filter(candidate__in=candidates)
            candidate_interview_status=InterviewCandidateStatus.objects.filter(interview__in=interviews)
            
            # Serialize the interview data
            interview_status_serializer = InterviewCandidateStatusCareerSerializer(candidate_interview_status, many=True)
            # Get the count of interviews
            interview_count = interviews.count()
            
            return Response({
                'data': interview_status_serializer.data,
                'interview': interview_count,
            })
            
        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


# class CoresignalCollectApi(APIView):
#     def get(self, request):
#         # Get the ID from the request parameters or use a default
#         search_view = SmartJobSearchAPIView()
#         q = (request.query_params.get('prompt') or request.query_params.get('q') or '').strip()
#         keyword_str = request.query_params.get('keyword')
        
#         # Parse keyword if it's a JSON string, otherwise use parse_query
#         if keyword_str:
#             try:
#                 import json
#                 keyword = json.loads(keyword_str)
#                 # keyword = keyword.get('jobTitles', '')
#                 if len(keyword.get('jobTitles', '')) == 0 or len(keyword.get('location', '')) == 0 or len(keyword.get('skills', '')) == 0 or len(keyword.get('education', '')) == 0 or keyword.get('experienceRange', '') == '':
#                     keyword = parse_query(q)        
#             except json.JSONDecodeError:
#                 keyword = parse_query(q)
#         else:
#             keyword = parse_query(q)

#         # print(parsed)

#         # url = "https://api.coresignal.com/cdapi/v2/employee_multi_source/search/es_dsl/preview?page=1"
#         url = f"https://api.coresignal.com/cdapi/v2/employee_multi_source/search/es_dsl/preview?page=1"
#         headers = {
#             "Content-Type": "application/json",
#             "apikey": "ZVDg8OGeOuvBk24DDYJvYfDqMBZTTzhq"
#         }

#         # Build the base query
#         payload = {
#             "query": {
#                 "bool": {
#                     "must": [],
#                     "should": [],
#                     "filter": []
#                 }
#             }
#         }

#         # Add search terms to the query
#         if keyword.get('jobTitles') or keyword.get('skills'):
#             must_queries = []
            
#             # Handle job titles search
#             if keyword.get('jobTitles'):
#                 for title in keyword['jobTitles']:
#                     must_queries.append({
#                         "multi_match": {
#                             "query": title,
#                             "fields": [
#                                 "headline^3",
#                                 "experience.title^3",
#                                 "experience.description^2",
#                                 "summary^2"
#                             ],
#                             "fuzziness": "AUTO"
#                         }
#                     })
            
#             # Handle skills search
#             if keyword.get('skills'):
#                 must_queries.append({
#                     "multi_match": {
#                         "query": ' '.join(keyword['skills']),
#                         "fields": [
#                             "skills.name^3",
#                             "experience.description^2",
#                             "summary^1"
#                         ],
#                         "operator": "or"
#                     }
#                 })
            
#             # Add all must queries to the main query
#             if must_queries:
#                 payload["query"]["bool"]["must"].extend(must_queries)

#         # Add experience filter if available
#         if keyword.get('experienceRange'):
#             try:
#                 # Parse experience range (e.g., "2 - 10" -> min=2, max=10)
#                 exp_range = [x.strip() for x in keyword['experienceRange'].split('-')]
#                 min_exp_years = float(exp_range[0]) if exp_range[0] else 0
#                 max_exp_years = float(exp_range[1]) if len(exp_range) > 1 else 30  # Default max 30 years if not specified
                
#                 # Convert years to months
#                 min_exp_months = int(min_exp_years * 12)
#                 max_exp_months = int(max_exp_years * 12)

#                 payload["query"]["bool"]["filter"].append({
#                     "range": {
#                         "total_experience_duration_months": {
#                             "gte": min_exp_months,
#                             "lte": max_exp_months
#                         }
#                     }
#                 })
#             except (ValueError, TypeError):
#                 pass  # Skip if experience is not a valid number

#         # Add location filter if available
#         if keyword.get('location') and keyword['location']:
#             for loc in keyword['location']:
#                 location_query = {
#                     "nested": {
#                         "path": "experience",
#                         "query": {
#                             "multi_match": {
#                                 "query": loc,
#                                 "fields": [
#                                     "experience.location",
#                                     "experience.location_country",
#                                     "experience.location_region"
#                                 ],
#                                 "fuzziness": "AUTO"
#                             }
#                         }
#                     }
#                 }
#                 payload["query"]["bool"]["should"].append(location_query)
#                 # Also search in current location
#                 payload["query"]["bool"]["should"].append({
#                     "multi_match": {
#                         "query": loc,
#                         "fields": [
#                             "location",
#                             "location_country",
#                             "location_region"
#                         ],
#                         "fuzziness": "AUTO"
#                     }
#                 })
#             # At least one location should match
#             if keyword['location']:  # Only set if there are locations
#                 payload["query"]["bool"]["minimum_should_match"] = 1

#         # If no specific filters, do a general search
#         if not any([payload["query"]["bool"]["must"], payload["query"]["bool"]["should"], payload["query"]["bool"]["filter"]]):
#             payload = {
#                 "query": {
#                     "match_all": {}
#                 },
#                 "sort": [
#                     {"_score": {"order": "desc"}},
#                     {"total_experience_duration_months": {"order": "desc"}}
#                 ],
#                 "size": 10
#             }
#         print("Final payload:", payload)

#         # print("payload",payload)
#         try:
#             response = requests.post(url, headers=headers, json=payload)
#             response.raise_for_status()
#         except Exception as e:
#             print("Error:", e)
#             return Response({"error": str("Your available credits have been used up. Kindly make a payment to retrieve candidate details")}, status=500)

#         if response.json():
#             for response_item in response.json():
#                 # First, try to get an existing record
#                 existing = CoresignalPreview.objects.filter(coresignal_id=response_item.get("id")).first()
                
#                 # Only create or update if no record exists or if existing record has is_list=True
#                 if not existing or existing.is_list:
#                     CoresignalPreview.objects.update_or_create(
#                         coresignal_id=response_item.get("id"),
#                         defaults={'data': response_item, 'is_list': True}
#                     )
#         return Response(response.json())

from .utils import parse_query
from .search_payload import build_es_payload
from .services import search_db, fetch_from_api
from .models import CoresignalPreview

PAGE_SIZE = 30
API_PAGE_SIZE = 20


class CoresignalCollectApi(APIView):
    def get(self, request):
        from account.models import Company

        page = int(request.query_params.get("page", 1))
        query = (request.query_params.get("prompt") or "").strip()
        keyword_str = request.query_params.get("keyword", "{}")
        keyword = json.loads(keyword_str) if keyword_str else {}

         # Get the current user's company
        try:
            company = Company.objects.get(id=request.user.company_id)
        except (Company.DoesNotExist, AttributeError):
            return Response({"error": "Company not found or user not authenticated"}, status=400)


        # If no valid search criteria in keyword, parse the query string
        if not any([
            keyword.get("education"),
            keyword.get("jobTitles"),
            keyword.get("location"),
            keyword.get("skills"),
            keyword.get("experienceRange")
        ]):
            keyword = parse_query(query)
        payload = build_es_payload(keyword)

        # -----------------------------
        # DB total
        # -----------------------------
        db_qs = CoresignalPreview.objects.filter(is_list=True).order_by("-created_at")
        db_total = db_qs.count()

        global_start = (page - 1) * PAGE_SIZE
        global_end = global_start + PAGE_SIZE

        results = []

        # -----------------------------
        # 1️⃣ TAKE FROM DB
        # -----------------------------
        if global_start < db_total:
            db_slice = db_qs[global_start:min(global_end, db_total)]
            # Get all candidate IDs for batch query
            candidate_ids = [str(obj.coresignal_id) for obj in db_slice]
            
            # Get all visited candidates for this company in a single query
            visited_candidates = set()
            try:
                visited_company = CoresignalCandidateVisiteCompany.objects.filter(
                    company=company,
                    visitescandidatelist__in=db_slice
                ).prefetch_related('visitescandidatelist').first()
                
                if visited_company:
                    visited_candidates = set(str(c.coresignal_id) for c in visited_company.visitescandidatelist.all())
            except CoresignalCandidateVisiteCompany.DoesNotExist:
                pass
            
            # Prepare results with isViewed status
            for obj in db_slice:
                candidate_data = obj.data
                candidate_data['isViewed'] = str(obj.coresignal_id) in visited_candidates
                candidate_data['isList'] = obj.is_list
                results.append(candidate_data)

        # If page is fully filled by DB → STOP
        if len(results) == PAGE_SIZE:
            return Response({
                "page": page,
                "source": "database",
                "cradite": 0 if page > 1 else 1,
                "results": results
            })

        # -----------------------------
        # 2️⃣ TAKE FROM API (ONLY REMAINING)
        # -----------------------------
        remaining = PAGE_SIZE - len(results)

        # Where API starts globally
        api_start_index = max(0, global_start - db_total)

        api_page = (api_start_index // API_PAGE_SIZE) + 1
        api_page_offset = api_start_index % API_PAGE_SIZE

        api_data = fetch_from_api(payload, api_page)
        print(len(api_data))
        
        # api_slice = api_data[
        #     api_page_offset : api_page_offset + remaining
        # ]
        for item in api_data:
            candidate, created = CoresignalPreview.objects.update_or_create(
                coresignal_id=item["id"],
                defaults={"data": item, "is_list": True}
            )
            # Check if this candidate was already visited by the company
            visited_candidates = set()
            try:
                visited_company = CoresignalCandidateVisiteCompany.objects.filter(
                    company=company,
                    visitescandidatelist=candidate
                ).first()
                
                if visited_company:
                    visited_candidates = set(str(c.coresignal_id) for c in visited_company.visitescandidatelist.all())
            except CoresignalCandidateVisiteCompany.DoesNotExist:
                pass
            
            # Add isViewed status to the candidate data
            item['isViewed'] = str(candidate.coresignal_id) in visited_candidates
            item['isList'] = candidate.is_list
            results.append(item)

        return Response({
            "page": page,
            "cradite": 0 if page > 1 else 1,
            "source": "database+api" if db_total > 0 else "api",
            "results": results
        })


class CoresignalPreviewApi(APIView):
    def get(self, request, id):
        from account.models import Company  # Import Company model

        company = Company.objects.get(id=request.user.company_id)
        
        try:
            core_signal = CoresignalPreview.objects.get(coresignal_id=id)
            if not core_signal.is_list:
                # Get the company instance using company_id
                obj, created = CoresignalCandidateVisiteCompany.objects.get_or_create(company=company)
                obj.visitescandidatelist.add(core_signal)
                if core_signal:
                    serializer = CoresignalPreviewSerializer(core_signal)
                    return Response(serializer.data)
        except CoresignalPreview.DoesNotExist:
            pass
        except Company.DoesNotExist:
            return Response({"error": "Company not found"}, status=404)
        

        url = f"https://api.coresignal.com/cdapi/v2/employee_multi_source/collect/{id}"
        
        # Set up headers with proper content type and API key
        headers = {
            "Content-Type": "application/json",
            "apikey": "ZVDg8OGeOuvBk24DDYJvYfDqMBZTTzhq"
        }
        
        try:
            # Make the POST request with headers
            response = requests.get(url, headers=headers)
            response.raise_for_status()  # Raise exception for HTTP errors
            
            # First try to get an existing record with is_list=False
            try:
                core_signal = CoresignalPreview.objects.get(coresignal_id=id, is_list=True)
                # Update the existing record with new data but keep is_list=False
                core_signal.data = response.json()
                core_signal.is_list =False
                core_signal.save()
                created = False
            except CoresignalPreview.DoesNotExist:
                # If no record with is_list=False exists, create a new one
                core_signal, created = CoresignalPreview.objects.get_or_create(
                    coresignal_id=id,
                    defaults={'data': response.json(), 'is_list': False}
                )
            obj, created = CoresignalCandidateVisiteCompany.objects.get_or_create(company=company)
            obj.visitescandidatelist.add(core_signal)
            serializer = CoresignalPreviewSerializer(core_signal)
            # Return the JSON response                
            return Response(serializer.data, status=status.HTTP_200_OK)
            
        except requests.exceptions.RequestException as e:
            # Return detailed error information
            error_details = {
                "error": str(e),
                "details": str(e.response.text) if hasattr(e, 'response') and e.response else "No response details"
            }
            status_code = e.response.status_code if hasattr(e, 'response') and e.response else 500
            return Response(error_details, status=status_code)

class CoresignalCandidateStatusView(APIView):
    """
    API View for CoresignalCandidateStatus model
    """

    def get(self, request, *args, **kwargs):
        """
        Get all CoresignalCandidateStatus records
        """
        statuses = CoresignalCandidateStatus.objects.all()
        data = CoresignalCandidateStatusSerializer(statuses, many=True).data
        return Response(data, status=status.HTTP_200_OK)

    def post(self, request, *args, **kwargs):
        """
        Create or update a CoresignalCandidateStatus record
        - If coresignal_candidate and company exist, update the status
        - If coresignal_candidate exists but company is different, create a new record
        """
        candidate_id = request.data.get('candidate_id')
        status_value = request.data.get('status')
        company_id = request.user.company_id

        try:
            try:
                coresignal_candidate = CoresignalPreview.objects.get(coresignal_id=candidate_id)
            except CoresignalPreview.DoesNotExist:            
                url = f"https://api.coresignal.com/cdapi/v2/employee_multi_source/collect/{candidate_id}"
                headers = {
                "Content-Type": "application/json",
                "apikey": "ZVDg8OGeOuvBk24DDYJvYfDqMBZTTzhq"
                }
                response = requests.get(url, headers=headers)
                response.raise_for_status()  # Raise exception for HTTP errors
                coresignal_candidate = CoresignalPreview.objects.create(coresignal_id=candidate_id, data=response.json())

            company = Company.objects.get(id=company_id)
            
            # Try to get existing record with same coresignal_candidate and company
            status_obj, created = CoresignalCandidateStatus.objects.get_or_create(
                coresignal_candidate=coresignal_candidate,
                company=company,
                defaults={'status': status_value}
            )
            
            # If record exists, update the status
            if not created:
                status_obj.status = status_value
                status_obj.save()
                message = 'CoresignalCandidateStatus updated successfully'
            else:
                message = 'CoresignalCandidateStatus created successfully'
                
            return Response({'status': 'success', 'message': message})
            
        except Exception as e:
            return Response(
                {'status': 'error', 'message': str(e)},
                status=status.HTTP_400_BAD_REQUEST
            )


def check_delivery_status(subject, body, from_email):
    bounce_senders = [
        'mailer-daemon',
        'postmaster'
    ]

    bounce_subject_keywords = [
        'delivery status notification',
        'undelivered mail',
        'delivery failed',
        'address not found'
    ]

    bounce_body_keywords = [
        'address not found',
        'does not exist',
        'recipient address rejected',
        '550 5.1.1',
        'delivery failed'
    ]

    # Normalize text
    subject_l = (subject or "").lower()
    body_l = (body or "").lower()
    from_l = (from_email or "").lower()

    # Check sender
    if any(sender in from_l for sender in bounce_senders):
        return 'FAILED'

    # Check subject
    if any(keyword in subject_l for keyword in bounce_subject_keywords):
        return 'FAILED'

    # Check body
    if any(keyword in body_l for keyword in bounce_body_keywords):
        return 'FAILED'

    return 'DELIVERED'



class GmailSentEmailView2:
    def fetch_sent_emails(self, request, target_email):
        try:
            from django.conf import settings
            
            # Use email credentials from Django settings
            sender_mail = settings.EMAIL_HOST_USER
            auth_password = settings.EMAIL_HOST_PASSWORD
            
            if not all([sender_mail, auth_password]):
                return {'code': 403, 'message': 'Email settings are not properly configured in the system.'}
            
            # Use Gmail's IMAP server for sent emails
            IMAP_SERVER = 'imap.gmail.com'
            IMAP_PORT = 993
            mail = imaplib.IMAP4_SSL(IMAP_SERVER, IMAP_PORT)
            mail.login(sender_mail, auth_password)
    
            email_list = []

            # Try to select the Sent Mail folder
            try:
                mail.select('"[Gmail]/Sent Mail"')
            except:
                try:
                    mail.select('"[Gmail]/Sent"')
                except:
                    return Response({"error": "Could not find Sent Mail folder"}, status=400)

            search_criteria = f'(TO "{target_email}")'
            result, data = mail.search(None, search_criteria)

            if result == 'OK':
                mail_ids = data[0].split()

                for num in mail_ids:
                    result, msg_data = mail.fetch(num, '(RFC822)')
                    
                    raw_email = msg_data[0][1]
                    msg = email.message_from_bytes(raw_email)
                    print('msg',msg)

                    subject = email.header.decode_header(msg['subject'])[0][0]
                    if isinstance(subject, bytes):
                        subject = subject.decode('utf-8', errors='replace')

                    from_email = email.header.decode_header(msg['from'])[0][0]
                    if isinstance(from_email, bytes):
                        from_email = from_email.decode('utf-8', errors='replace')

                    to_email = email.header.decode_header(msg['to'])[0][0]
                    if isinstance(to_email, bytes):
                        to_email = to_email.decode('utf-8', errors='replace')
                        
                    date = email.header.decode_header(msg['date'])[0][0]
                    if isinstance(date, bytes):
                        date = date.decode('utf-8', errors='replace')

                    body = ""
                    if msg.is_multipart():
                        for part in msg.walk():
                            if part.get_content_type() == "text/plain":
                                body = part.get_payload(decode=True).decode('utf-8', errors='replace')
                                break
                    else:
                        body = msg.get_payload(decode=True).decode('utf-8', errors='replace')
                    
                    delivery_status = check_delivery_status(subject, body, from_email)
                    message_id = msg.get('Message-ID', '').strip()



                    email_list.append({
                        'subject': subject,
                        'from': from_email,
                        'to': to_email,
                        'date': date,
                        'body': body,
                        'delivery_status': delivery_status,  # ✅ NEW'
                        "message_id":message_id
                    })

            mail.logout()  # Logout from IMAP server
            return Response(email_list)

        except imaplib.IMAP4.error as e:
            return {"error": f"IMAP error: {str(e)}"}
        except Exception as e:
            return {"error": str(e)}


# class GmailAllEmailsView(APIView):
#     def get(self, request, id):
#         try:
#             candidate = Candidate.objects.get(id=id)
#             target_email = candidate.webform_candidate_data['Personal Details']['email']

#             # Create instances of both views
#             receive_email_view = GmailReceiveEmailView2()
#             sent_email_view = GmailSentEmailView2()

#             # Use ThreadPoolExecutor to execute both methods concurrently
#             with ThreadPoolExecutor(max_workers=8) as executor:
#                 future_receive = executor.submit(receive_email_view.fetch_received_emails, request, target_email)
#                 future_sent = executor.submit(sent_email_view.fetch_sent_emails, request, target_email)

#                 # Wait for the results and get the actual data
#                 received_emails = future_receive.result()
#                 sent_emails = future_sent.result()

#             # Extract the data from the Response objects
#             if isinstance(received_emails, Response):
#                 received_emails = received_emails.data
#             if isinstance(sent_emails, Response):
#                 sent_emails = sent_emails.data
                
#             if not received_emails:
#                 received_emails = None
#             if not sent_emails:
#                 sent_emails = None

#             # Construct the response data
#             response_data = {
#                 'received_emails': received_emails,
#                 'sent_emails': sent_emails
#             }

#             return JsonResponse(response_data)

#         except Candidate.DoesNotExist:
#             return JsonResponse({"error": "Candidate not found."}, status=404)
#         except Exception as e:
#             return JsonResponse({"error": str(e)}, status=400)


class GmailReceiveEmailView2:
    def fetch_received_emails(self, request, target_email):
        try:
            from django.conf import settings
            
            # Use email credentials from Django settings
            sender_mail = settings.EMAIL_HOST_USER
            auth_password = settings.EMAIL_HOST_PASSWORD
            
            if not all([sender_mail, auth_password]):
                return {'code': 403, 'message': 'Email settings are not properly configured in the system.'}
            
            # Use Gmail's IMAP server for receiving emails
            IMAP_SERVER = 'imap.gmail.com'
            IMAP_PORT = 993
            mail = imaplib.IMAP4_SSL(IMAP_SERVER, IMAP_PORT)
            mail.login(sender_mail, auth_password)
            mail.select('inbox')  # Select the inbox for received emails

            search_criteria = f'(OR FROM "{target_email}" TO "{target_email}")'
            result, data = mail.search(None, search_criteria)

            email_list = []  # Initialize email list

            if result == 'OK':
                # Get the list of email IDs
                mail_ids = data[0].split()
                
                if not mail_ids:
                    # Return an empty list or None if no emails are found
                    return Response(None)

                # Decode each mail ID from bytes to string
                mail_ids = [id.decode('utf-8') for id in mail_ids]

                # Fetch emails in a batch instead of one by one
                email_id_string = ','.join(mail_ids)  # Join the email IDs as a string for fetching
                result, msg_data = mail.fetch(email_id_string, '(RFC822)')

                if result == 'OK':
                    for response_part in msg_data:
                        if isinstance(response_part, tuple):
                            raw_email = response_part[1]
                            msg = email.message_from_bytes(raw_email)

                            # Decode the email subject
                            subject = email.header.decode_header(msg['subject'])[0][0]
                            if isinstance(subject, bytes):
                                subject = subject.decode('utf-8', errors='replace')

                            # Decode the sender email
                            from_email = email.header.decode_header(msg['from'])[0][0]
                            if isinstance(from_email, bytes):
                                from_email = from_email.decode('utf-8', errors='replace')

                            to_email = email.header.decode_header(msg['to'])[0][0]
                            if isinstance(to_email, bytes):
                                to_email = to_email.decode('utf-8', errors='replace')
                                
                            date = email.header.decode_header(msg['date'])[0][0]
                            if isinstance(date, bytes):
                                date = date.decode('utf-8', errors='replace')

                            # Extract the email body (plain text or first part)
                            body = ""
                            if msg.is_multipart():
                                for part in msg.walk():
                                    content_type = part.get_content_type()
                                    content_disposition = str(part.get("Content-Disposition"))

                                    if content_type == "text/plain" and "attachment" not in content_disposition:
                                        body = part.get_payload(decode=True).decode('utf-8', errors='replace')
                                        break
                            else:
                                body = msg.get_payload(decode=True).decode('utf-8', errors='replace')

                            message_id = msg.get('Message-ID', '').strip()

                            # Append the filtered email information
                            email_list.append({
                                'subject': subject,
                                'from': from_email,
                                'to': to_email,
                                'date': date,
                                'body': body,  # Limit body length for performance (first 200 chars)
                                'Message-ID': message_id
                            })
                
            mail.logout()  # Logout from IMAP server
            return Response(email_list if email_list else [])

        except imaplib.IMAP4.error as e:
            return {"error": f"IMAP error: {str(e)}"}
        except Exception as e:
            return {"error": str(e)}


class CoresignalCandidateSEmailView(APIView):
    
    def get(self, request, id: int = None) -> JsonResponse:
        """
        Fetch received and sent emails for a specific candidate.
        
        Args:
            request: Django request object
            candidate_id: ID of the candidate to fetch emails for
            
        Returns:
            JsonResponse containing received_emails and sent_emails
        """
        try:
            # Get candidate emails
            coresignal_preview = get_object_or_404(CoresignalPreview, coresignal_id=id)
            emails = coresignal_preview.data.get('email')
            # emails = ["1010@gmail.com"]
            
            if not emails:
                return JsonResponse({"error": "No email found for this candidate"}, status=400)

            # Create instances of both views
            receive_email_view = GmailReceiveEmailView2()
            sent_email_view = GmailSentEmailView2()

            # Initialize empty lists to store results
            all_received_emails = []
            all_sent_emails = []

            # Process each email
            for email in emails:
                # Use ThreadPoolExecutor to execute both methods concurrently for each email
                with ThreadPoolExecutor(max_workers=2) as executor:
                    future_receive = executor.submit(receive_email_view.fetch_received_emails, request, email)
                    future_sent = executor.submit(sent_email_view.fetch_sent_emails, request, email)

                    # Wait for the results and get the actual data
                    received_emails = future_receive.result()
                    sent_emails = future_sent.result()

                    # Add to our combined results if we got valid data
                    # Handle received_emails
                    if received_emails and not isinstance(received_emails, dict) or (isinstance(received_emails, dict) and 'error' not in received_emails):
                        emails_to_add = received_emails.data if hasattr(received_emails, 'data') else (received_emails if isinstance(received_emails, list) else [])
                        if emails_to_add:
                            all_received_emails.extend(emails_to_add)
                    
                    # Handle sent_emails
                    if sent_emails and not isinstance(sent_emails, dict) or (isinstance(sent_emails, dict) and 'error' not in sent_emails):
                        emails_to_add = sent_emails.data if hasattr(sent_emails, 'data') else (sent_emails if isinstance(sent_emails, list) else [])
                        if emails_to_add:
                            all_sent_emails.extend(emails_to_add)

            # Construct the response data with combined results
            response_data = {
                'received_emails': all_received_emails if all_received_emails else [],
                'send_emails': all_sent_emails if all_sent_emails else []
            }

            return JsonResponse(response_data, safe=False)

        except CoresignalPreview.DoesNotExist:
            return JsonResponse({"error": "Candidate not found."}, status=404)
        except Exception as e:
            logger.error(f"Error in CoresignalCandidateSEmailSendView: {str(e)}", exc_info=True)
            return JsonResponse({"error": "An error occurred while processing your request."}, status=500)

    def post(self, request, *args, **kwargs):
        data = request.data
        subject = data.get('subject')
        body = data.get('body')
        candidate_id = data.get('candidate_id')
        message_id =data.get('message_id')
        
        coresignal_preview = get_object_or_404(CoresignalPreview, coresignal_id=candidate_id)
        email = coresignal_preview.data.get('email')
        # print(email)
        if not all([subject, body]):
            return Response(
                {'error': 'Subject and body are required'},
                status=status.HTTP_400_BAD_REQUEST
            )
            
        try:
            from django.core.mail import EmailMessage
            
            email = EmailMessage(
                subject=subject,
                body=body,
                from_email=None,  # Uses DEFAULT_FROM_EMAIL from settings
                to=[email],
                reply_to=message_id
            )
            email.send(fail_silently=False)
            
            return Response(
                {'message': 'Email sent successfully'},
                status=status.HTTP_200_OK
            )
        except Exception as e:
            return Response(
                {'error': f'Failed to send email: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class EmailReplyGeneratorAPIView(APIView):
    """
    Generate email subject and message using past email content
    """
    @check_subscription_and_credits_for_ai(feature_code = "AI_CREDITS", usage_code="AI_Outreach_Email_Follow_ups")
    @handle_ai_credits(feature_code="AI_CREDITS", usage_code="AI_Outreach_Email_Follow_ups")
    def post(self, request):
        past_email = request.data.get("email")

        candidate_name = request.data.get("candidate_name")
        job_title = request.data.get("job_title")
        company_name = request.user.company_id
        company_name = Company.objects.get(id=company_name)
        sender_name = f'{request.user.first_name} {request.user.last_name}'

        client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)

        # ---------- CASE 1: Past email exists ----------
        if past_email:
            prompt = f"""
            Based on the following email conversation, generate:
            1. A clear and professional email subject
            2. A polite and professional email reply message

            Email Content:
            {past_email}

            Return the result in the following format:
            Subject: <subject here>
            Message: <message here>
            """

        # ---------- CASE 2: No past email (Fresh email) ----------
        else:
            if not all([candidate_name, job_title, company_name, sender_name]):
                return Response(
                    {
                        "error": "When past email is not provided, candidate_name, job_title, company_name, and sender_name are required."
                    },
                    status=status.HTTP_400_BAD_REQUEST
                )

            prompt = f"""
            Write a professional job-related email with the following details:

            Candidate Name: {candidate_name}
            Job Title: {job_title}
            Company Name: {company_name}
            Sender Name: {sender_name}

            The email should be polite, professional, and suitable for formal communication.

            Return the result in the following format:
            Subject: <subject here>
            Message: <message here>
            """

        try:
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a professional email writing assistant."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.5,
                max_tokens=300
            )

            input_tokens = response.usage.prompt_tokens
            output_tokens = response.usage.completion_tokens

            # 🔹 Cost (update if pricing changes)
            input_cost = (input_tokens / 1_000_000) * 5.00
            output_cost = (output_tokens / 1_000_000) * 15.00

            print(
                f"Email Generate → in:{input_tokens} out:{output_tokens} "
                f"cost:${input_cost + output_cost:.6f}"
            )
            text = response.choices[0].message.content

            subject = ""
            message = ""

            if "Subject:" in text and "Message:" in text:
                subject = text.split("Subject:")[1].split("Message:")[0].strip()
                message = text.split("Message:")[1].strip()
            else:
                message = text

            return Response(
                {
                    "subject": subject,
                    "message": message
                },
                status=status.HTTP_200_OK
            )

        except Exception as e:
            return Response(
                {"error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

class CandidateVisiteCompanyListView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def get(self, request):
        try:
            # Get the company's visited candidates
            visited_company = get_object_or_404(
                CoresignalCandidateVisiteCompany,
                company=request.user.company_id
            )
            
            # Get all visited candidates for this company
            visited_candidates = visited_company.visitescandidatelist.all()
            
            # Serialize the data
            serializer = CoresignalPreviewSerializer(visited_candidates, many=True)
            
            return Response(serializer.data)
            
        except Exception as e:
            return Response({
                'success': False,
                'error': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)


def get_ai_matching_candidates(self, job_title, skills, description):
    """
    Find relevant candidates using OpenAI based on job title, skills, and description.
    Returns a list of candidate IDs sorted by relevance.
    """
    try:
        # Get all candidates with their details
        candidates = Candidate.objects.filter(company=self.request.user.company)
        
        # Prepare candidate data for AI analysis
        candidate_data = []
        for candidate in candidates:
            candidate_info = {
                'id': candidate.id,
                'name': f"{candidate.first_name or ''} {candidate.last_name or ''}".strip(),
                'skills': candidate.skills or '',
                'experience': f"{candidate.exp_years or 0} years" if candidate.exp_years else 'No experience',
                'qualification': candidate.highest_qualification or '',
                'current_job': candidate.cur_job or '',
                'summary': candidate.summary or ''
            }
            candidate_data.append(candidate_info)
        
        if not candidate_data:
            return []
        
        # Create a prompt for OpenAI
        prompt = f"""I need to find the most suitable candidates for the following job:
        
        Job Title: {job_title}
        Required Skills: {skills}
        Job Description: {description}
        
        Here are the candidates to evaluate. Please return a JSON array of candidate IDs sorted by relevance (most relevant first), 
        along with a relevance score between 0 and 1 (1 being the best match). Consider their skills, experience, and overall profile:
        
        {json.dumps(candidate_data, indent=2)}
        
        Return format: 
        {{
            "candidates": [
                {{"id": 123, "relevance_score": 0.95, "match_reason": "Skills and experience closely match the job requirements"}},
                ...
            ]
        }}
        """
        
        # Call OpenAI
        response = client.chat.completions.create(
            model="gpt-4o-mini-2024-07-18",
            messages=[
                {"role": "system", "content": "You are an AI assistant that helps match candidates to job descriptions. Analyze the candidates and return the most relevant ones based on the job requirements."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=2000
        )
        
        # Parse the response
        input_tokens = response.usage.prompt_tokens
        output_tokens = response.usage.completion_tokens

        # 🔹 Cost (update if pricing changes)
        input_cost = (input_tokens / 1_000_000) * 5.00
        output_cost = (output_tokens / 1_000_000) * 15.00

        print(
            f"Job Keyword Extract → in:{input_tokens} out:{output_tokens} "
            f"cost:${input_cost + output_cost:.6f}"
        )
        result = json.loads(response.choices[0].message.content)
        return result.get('candidates', [])
        
    except Exception as e:
        logger.error(f"Error in AI candidate matching: {str(e)}")
        return []

class Corsignalcandidatesearch(APIView):
    def get(self, request, id):
        try:
            job = get_object_or_404(Job, id=id)

            if not hasattr(job, 'dynamic_job_data'):
                return Response(
                    {'error': 'Job data not found'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            create_job_data = job.dynamic_job_data.get('Create Job', {})
            description_info = job.dynamic_job_data.get('Description Information', {})

            jobTitles = create_job_data.get('title', '')
            skills = create_job_data.get('speciality', '')
            description = description_info.get('description', '')
            experienceRange = create_job_data.get('exp_min', '') + ' - ' + create_job_data.get('exp_max', '') if create_job_data.get('exp_max', '') else create_job_data.get('exp_min', '')
            # experience_max = create_job_data.get('exp_max', '')

            if not jobTitles:
                return Response(
                    {'error': 'Job title is missing'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # OpenAI setup
            client = OpenAI(api_key=settings.OPENAI_API_KEY)

            # prompt_instruction = f"""
            # Create a 1-line candidate search query string that includes:
            # - Job title
            # - Key skills
            # - Experience (use a realistic number between 5–10 years)
            # - Optional location only if mentioned

            # The output must be plain text and parsable like:
            # "<job title> <skills> <years> years in <location>"

            # Job Title: {job_title}
            # Skills: {skill}
            # Experience Range: {experience_min} - {experience_max} years
            # description: {description}
            # """
            # prompt_instruction = f"""
            #     Create a single-line candidate search query string.

            #     Rules:
            #     - Include job title and key skills
            #     - Use ONLY ONE experience number selected from this range: {experience_min} to {experience_max} years
            #     - write a range like "5-10 years"
            #     - Use plain text only (no bullets, no quotes)
            #     - Add location ONLY if clearly mentioned in the description
            #     - The output must be parsable in the format:
            #     <job title> <skills> <years> years in <location>

            #     Input Data:
            #     Job Title: {job_title}
            #     Skills: {skill}
            #     Experience Range: {experience_min} - {experience_max} years
            #     Description: {description}
            #     """

            prompt_instruction = f"""
                You are an expert recruitment search assistant.

                Task:
                Generate a SINGLE valid JSON object.

                STRICT OUTPUT RULES:
                - Output ONLY valid JSON
                - No explanations, no markdown, no extra text
                - Keys and structure must match EXACTLY

                Required JSON Format:
                {{
                "promant": "<single-line candidate search query>",
                "keyword": {{
                    "education": [],
                    "experienceRange": "",
                    "jobTitles": [],
                    "location": [],
                    "skills": []
                }}
                }}

                Rules for "promant":
                - Single-line plain text
                - Format:
                <job title> <skills> <years> years in <location>
                - Choose ONE primary job title
                - Include only relevant skills
                - Convert experienceRange into a readable range like "5-10 years"
                - Include location ONLY if provided

                Rules for "keyword":
                - Fill values ONLY from input data
                - Do NOT invent values
                - If a field is missing, return empty array [] or empty string ""

                Input Data:
                Job Titles: {jobTitles}
                Skills: {skills}
                Experience Range: {experienceRange}
                Description: {description}
                """
                # Location: {location}
                # Education: {education}




            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You generate concise job search queries for candidate matching."},
                    {"role": "user", "content": prompt_instruction}
                ],
                # temperature=0.3,
                # max_tokens=50
            )

            input_tokens = response.usage.prompt_tokens
            output_tokens = response.usage.completion_tokens

            # 🔹 Cost (update if pricing changes)
            input_cost = (input_tokens / 1_000_000) * 5.00
            output_cost = (output_tokens / 1_000_000) * 15.00

            print(
                f"Job Keyword Extract → in:{input_tokens} out:{output_tokens} "
                f"cost:${input_cost + output_cost:.6f}"
            )


            raw_response = response.choices[0].message.content.strip()
            parsed_response = json.loads(raw_response)

            return Response({
                # "job_title": jobTitles,
                # "skill": skills,
                "generated_prompt": parsed_response,
                # "model": response.model,
                # "promant":response,
                # "usage": {
                #     "prompt_tokens": response.usage.prompt_tokens,
                #     "completion_tokens": response.usage.completion_tokens,
                #     "total_tokens": response.usage.total_tokens
                # }

            })

        except Exception as e:
            logger.error(f"Error in CoresignalCandidateSearch: {str(e)}")
            return Response(
                {'error': 'An error occurred while generating search prompt'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class CorsignalCandidateScrongin(APIView):
    def get(self, request,job_id,corsignal_id):
        
        job = get_object_or_404(Job, id = job_id)
        coresignal = get_object_or_404(CoresignalPreview, coresignal_id = corsignal_id)
        
        prompt = f"""You are an AI recruitment assistant that analyzes job-candidate compatibility.

        TASK:
        Analyze the following job posting and candidate profile, then provide a compatibility assessment.

        RULES:
        1. Prioritize job title, core responsibilities, and required skills
        2. Missing critical skills should significantly reduce the match score
        3. Partial or transferable skills should moderately reduce the score
        4. Consider experience and education as supporting factors
        5. Be more stringent when evaluating senior roles

        IMPORTANT: Your response MUST be valid JSON. Do not include any text outside the JSON structure.

        MATCHING GUIDELINES:
        - 0-39: Not Suitable (significant gaps in key requirements)
        - 40-69: Partially Suitable (some matches but with notable gaps)
        - 70-100: Suitable (good to excellent match with most requirements met)

        SCORING CRITERIA:
        1. Best Match (70-100): Candidate meets or exceeds most requirements
        - Has all critical skills
        - Experience level matches or exceeds requirements
        - Education and other factors are a good fit
        - person_match: true

        2. Good Match (40-69): Candidate meets some requirements
        - Has some but not all critical skills
        - Experience may be slightly below requirements
        - May have some gaps in qualifications
        - person_match: false

        3. Not Suitable (0-39): Significant gaps in requirements
        - Missing critical skills
        - Experience significantly below requirements
        - Major gaps in qualifications
        - person_match: false

        JOB DATA:
        {json.dumps(job.dynamic_job_data, indent=2) if job.dynamic_job_data else 'No job data available'}

        CANDIDATE DATA:
        {json.dumps(coresignal.data, indent=2) if coresignal.data else 'No candidate data available'}

        REQUIRED RESPONSE FORMAT (strict JSON, no markdown code blocks):
        {{
            "candidate_name": "string (candidate's full name)",
            "job_title": "string (job title being applied for)",
            "person_match": "boolean (true if candidate is a good fit)",
            "overall_match_percentage": "number (0-100)",
            "match_level": "string (best/good/partial/not)",
            "strengths": ["string"],
            "gaps": ["string"],
            "final_reason": "string (brief explanation of the match decision)"
        }}

        Note: All fields are required. For arrays, use empty arrays if no items.
        """

        try:
            # Make the API call
            response = client.chat.completions.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[
                    {"role": "system", "content": "You are an AI assistant that helps match candidates to job descriptions. Analyze the candidates and return the most relevant ones based on the job requirements."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=2000
            )
            
            # Calculate and log the cost
            if hasattr(response, 'usage') and response.usage:
                cost_data = calculate_cost("gpt-4o-mini-2024-07-18", response.usage)
                if cost_data:
                    print(
                        f"OpenAI API Cost - Model: {cost_data['model']}, "
                        f"Input Tokens: {cost_data['input_tokens']} (${cost_data['input_cost']:.6f}), "
                        f"Output Tokens: {cost_data['output_tokens']} (${cost_data['output_cost']:.6f}), "
                        f"Total Cost: ${cost_data['total_cost']:.6f} USD, "
                        f"Job ID: {job_id}, Coresignal ID: {corsignal_id}"
                    )
            
            # Parse the response
            result = json.loads(response.choices[0].message.content)
            # coresignal.scrongin = result
            # coresignal.save()
            return Response({"result":result,"job_id":job_id,"corsignal_id":corsignal_id,"is_list":coresignal.is_list}, status=status.HTTP_200_OK)
        except json.JSONDecodeError as e:
            return Response(
                {"error": "Failed to parse AI response", "details": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        except Exception as e:
            return Response(
                {"error": "An error occurred", "details": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        



class NotificationListView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        notifications = Notification.objects.filter(user=request.user).order_by('-created_at')[:20]
        serializer = NotificationSerializer(notifications, many=True)
        return Response(serializer.data)

class MarkNotificationReadView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request, pk):
        try:
            notification = Notification.objects.get(pk=pk, user=request.user)
            notification.is_read = True
            notification.save()
            return Response({"message": "Notification marked as read"})
        except Notification.DoesNotExist:
            return Response({"error": "Notification not found"}, status=status.HTTP_404_NOT_FOUND)
